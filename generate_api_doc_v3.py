#!/usr/bin/env python3
"""
Generate Basket4Me PWA - Van Sales API Reference V3 (Word Document)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json
import os

# ── Helpers ──────────────────────────────────────────────────────────────────

BASE_URL = "https://devassociate.frappe.cloud"
API_PREFIX = "/api/method/basket4me_pwa.api."
AUTH_PREFIX = "/api/method/basket4me_pwa.auth."

doc = Document()

# Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f"Heading {i}"]
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, font_name=None, font_size=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size

def add_header_row(table, headers, color="2E75B6"):
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, color)
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), font_size=Pt(9))

def make_table(headers, rows, header_color="2E75B6"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(table, headers, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_text(cell, str(val), font_size=Pt(9))
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F7FB")
    return table

def add_endpoint_summary_table(endpoints):
    rows = []
    for i, ep in enumerate(endpoints, 1):
        rows.append([str(i), ep["path"], ep["method"], ep["desc"]])
    make_table(["#", "Endpoint Path", "Method", "Description"], rows)

def add_json_block(title, data):
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    json_str = json.dumps(data, indent=2)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(json_str)
    run2.font.name = "Consolas"
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_param_table(params):
    rows = []
    for p in params:
        rows.append([p["name"], p["type"], p.get("required", "Yes"), p.get("desc", "")])
    make_table(["Parameter", "Type", "Required", "Description"], rows, header_color="2D8B4E")

def add_endpoint_detail(path, method, description, params=None, request_body=None, response_body=None):
    doc.add_paragraph()
    p = doc.add_paragraph()
    run_m = p.add_run(f"  {method}  ")
    run_m.bold = True
    run_m.font.size = Pt(10)
    if method in ("POST", "DELETE"):
        run_m.font.color.rgb = RGBColor(0xE0, 0x6C, 0x00)
    else:
        run_m.font.color.rgb = RGBColor(0x2D, 0x8B, 0x4E)
    run_url = p.add_run(f"  {BASE_URL}{path}")
    run_url.font.name = "Consolas"
    run_url.font.size = Pt(9)
    run_url.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p2 = doc.add_paragraph(description)
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].italic = True

    if params:
        doc.add_paragraph("Parameters:", style="List Bullet")
        add_param_table(params)

    if request_body:
        add_json_block("Request Body:", request_body)

    if response_body:
        add_json_block("Response:", response_body)

    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("_" * 90)
    run_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run_sep.font.size = Pt(6)

def section_heading(num, title):
    doc.add_page_break()
    doc.add_heading(f"Section {num}: {title}", level=1)

# ── TITLE PAGE ───────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title_p.add_run("Basket4Me PWA")
run_t.bold = True
run_t.font.size = Pt(32)
run_t.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = sub_p.add_run("Van Sales Mobile App\nAPI Reference V3")
run_s.font.size = Pt(22)
run_s.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ("Base URL: ", BASE_URL),
    ("\nVersion: ", "3.0"),
    ("\nDate: ", "March 2026"),
    ("\nAuthorization: ", "Token-based (api_key:api_secret)"),
]:
    r_l = info_p.add_run(label)
    r_l.bold = True
    r_l.font.size = Pt(12)
    r_l.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r_v = info_p.add_run(value)
    r_v.font.size = Pt(12)
    r_v.font.name = "Consolas"
    r_v.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── TABLE OF CONTENTS ────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1. Authentication (3 endpoints)",
    "2. Dashboard (3 endpoints)",
    "3. Customer Management (11 endpoints)",
    "4. Customer Route (1 endpoint)",
    "5. Items (8 endpoints)",
    "6. Price List (5 endpoints)",
    "7. Sales Order (8 endpoints)",
    "8. Delivery Note (10 endpoints) [Optional]",
    "9. Sales Invoice (10 endpoints)",
    "10. Sales Return (4 endpoints)",
    "11. Receipts / Payments (5 endpoints)",
    "12. Material Request (6 endpoints)",
    "13. Print (5 endpoints)",
    "14. Config & Settings (10 endpoints)",
    "15. Shift Management (3 endpoints)",
    "",
    "Appendix A: Auth Header Format",
    "Appendix B: Standard Error Response Format",
    "Appendix C: HTTP Status Codes",
    "Appendix D: Workflow Diagrams",
]
for item in toc_items:
    if item == "":
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(11)

# ── AUTH HEADERS SECTION ─────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading("Authentication & Authorization", level=1)

doc.add_paragraph(
    "The Basket4Me PWA API uses token-based authentication. "
    "Call the Login endpoint to obtain api_key and api_secret, "
    "then include them in all subsequent requests."
)

doc.add_heading("Token Authentication", level=2)
doc.add_paragraph(
    "1. Call the Login endpoint with usr and pwd.\n"
    "2. The response contains api_key and api_secret.\n"
    "3. Store both values securely in local storage.\n"
    "4. Include in ALL subsequent requests as a single header."
)

make_table(
    ["Header", "Format", "Example"],
    [
        ["Authorization", "token <api_key>:<api_secret>", "Authorization: token c12a33bb670b1e1:80c7720caff3272"],
    ],
)

doc.add_paragraph()
p_note = doc.add_paragraph()
run_note = p_note.add_run("Important: ")
run_note.bold = True
p_note.add_run(
    "Token auth works for both GET and POST requests. No Cookie or CSRF token headers needed. "
    "This avoids the browser 'Refused to set unsafe header Cookie' error in PWA/frontend apps."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: Authentication
# ══════════════════════════════════════════════════════════════════════════════

section_heading(1, "Authentication")

auth_endpoints = [
    {"path": AUTH_PREFIX + "user_login", "method": "POST", "desc": "Login and obtain API credentials"},
    {"path": AUTH_PREFIX + "logout", "method": "POST", "desc": "Logout and invalidate session"},
    {"path": AUTH_PREFIX + "get_user_details", "method": "POST", "desc": "Get current user details and permissions"},
]
add_endpoint_summary_table(auth_endpoints)

# 1.1 user_login
add_endpoint_detail(
    AUTH_PREFIX + "user_login", "POST",
    "Authenticate user and return API credentials. Allows guest access.",
    params=[
        {"name": "usr", "type": "String", "required": "Yes", "desc": "Username or email address"},
        {"name": "pwd", "type": "String", "required": "Yes", "desc": "User password"},
        {"name": "device_id", "type": "String", "required": "No", "desc": "Mobile device identifier"},
    ],
    request_body={"usr": "salesman@example.com", "pwd": "password123", "device_id": "ABC-123-DEF"},
    response_body={
        "message": {
            "sid": "a1b2c3d4e5f6",
            "csrf_token": "xyz789token",
            "api_key": "c12a33bb670b1e1",
            "api_secret": "80c7720caff3272",
            "username": "salesman@example.com",
            "email": "salesman@example.com",
            "mobile_no": "+966501234567"
        }
    },
)

# 1.2 logout
add_endpoint_detail(
    AUTH_PREFIX + "logout", "POST",
    "Logout user and invalidate the current session. Allows guest access.",
    params=[
        {"name": "usr", "type": "String", "required": "Yes", "desc": "Username or email of the user to logout"},
    ],
    request_body={"usr": "salesman@example.com"},
    response_body={"message": "Logged Out"},
)

# 1.3 get_user_details
add_endpoint_detail(
    AUTH_PREFIX + "get_user_details", "POST",
    "Retrieve detailed user information including company details and permissions.",
    params=[
        {"name": "sid", "type": "String", "required": "No", "desc": "Session ID"},
        {"name": "user_id", "type": "String", "required": "No", "desc": "User ID or email"},
    ],
    request_body={"sid": "a1b2c3d4e5f6", "user_id": "salesman@example.com"},
    response_body={
        "message": {
            "sid": "a1b2c3d4e5f6",
            "api_key": "c12a33bb670b1e1",
            "api_secret": "80c7720caff3272",
            "username": "salesman@example.com",
            "email": "salesman@example.com",
            "full_name": "Ahmed Al-Saleh",
            "mobile_no": "+966501234567",
            "user_image": "/files/user_avatar.png",
            "company_name": "Al-Marai Distribution LLC",
            "company_logo": "/files/company_logo.png",
            "employee_id": "HR-EMP-00042",
            "permissions": {
                "Sales Order": {"create": 1, "read": 1, "write": 1, "delete": 0, "submit": 1, "cancel": 1},
                "Delivery Note": {"create": 1, "read": 1, "write": 1, "delete": 0, "submit": 1, "cancel": 1},
                "Sales Invoice": {"create": 1, "read": 1, "write": 1, "delete": 0, "submit": 1, "cancel": 1},
                "Sales Return": {"create": 1, "read": 1, "write": 0, "delete": 0, "submit": 0, "cancel": 0},
                "Payment Entry": {"create": 1, "read": 1, "write": 1, "delete": 0, "submit": 1, "cancel": 1},
            }
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: Dashboard
# ══════════════════════════════════════════════════════════════════════════════

section_heading(2, "Dashboard")

dash_endpoints = [
    {"path": API_PREFIX + "get_dashboard_summary", "method": "GET", "desc": "Get summary metrics for the dashboard"},
    {"path": API_PREFIX + "get_sales_metrics", "method": "GET", "desc": "Get sales metrics with filters"},
    {"path": API_PREFIX + "get_daily_sales_report", "method": "GET", "desc": "Get daily sales report"},
]
add_endpoint_summary_table(dash_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_dashboard_summary", "GET",
    "Retrieve dashboard summary metrics filtered by time period.",
    params=[
        {"name": "period", "type": "String", "required": "No", "desc": "Period filter: 'daily', 'weekly', or 'monthly'"},
        {"name": "from_date", "type": "String", "required": "No", "desc": "Start date (YYYY-MM-DD)"},
        {"name": "to_date", "type": "String", "required": "No", "desc": "End date (YYYY-MM-DD)"},
    ],
    response_body={
        "message": {
            "total_sales": 125000.00,
            "total_orders": 45,
            "total_customers_visited": 32,
            "total_collections": 98000.00,
            "outstanding_amount": 27000.00,
            "total_returns": 3500.00
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_sales_metrics", "GET",
    "Get detailed sales metrics with date range and sales person filter.",
    params=[
        {"name": "from_date", "type": "String", "required": "No", "desc": "Start date (YYYY-MM-DD)"},
        {"name": "to_date", "type": "String", "required": "No", "desc": "End date (YYYY-MM-DD)"},
        {"name": "sales_person_filter", "type": "String", "required": "No", "desc": "Filter by sales person name"},
    ],
    response_body={
        "message": {
            "total_invoiced": 95000.00,
            "total_collected": 82000.00,
            "total_outstanding": 13000.00,
            "invoice_count": 28,
            "avg_invoice_value": 3392.86
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_daily_sales_report", "GET",
    "Get the daily sales report for the current user. No parameters required.",
    response_body={
        "message": {
            "date": "2026-03-21",
            "salesperson": "Ahmed Al-Saleh",
            "orders_created": 12,
            "invoices_created": 10,
            "total_sales": 45000.00,
            "total_collections": 38000.00,
            "customers_visited": 15
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: Customer Management
# ══════════════════════════════════════════════════════════════════════════════

section_heading(3, "Customer Management")

cust_endpoints = [
    {"path": API_PREFIX + "get_customer_list", "method": "GET", "desc": "Get customer list (basic)"},
    {"path": API_PREFIX + "get_customer_list_v2", "method": "GET", "desc": "Get customer list with pagination & filters"},
    {"path": API_PREFIX + "get_customer_list_with_effective_price_list", "method": "GET", "desc": "Get customers with effective price list"},
    {"path": API_PREFIX + "get_customer_detail", "method": "POST", "desc": "Get customer details"},
    {"path": API_PREFIX + "create_customer", "method": "POST", "desc": "Create a new customer"},
    {"path": API_PREFIX + "get_customer_invoices", "method": "GET", "desc": "Get customer invoices"},
    {"path": API_PREFIX + "customer_invoice_details", "method": "POST", "desc": "Get customer invoice details"},
    {"path": API_PREFIX + "get_customer_invoice_aging", "method": "GET", "desc": "Get customer invoice aging report"},
    {"path": API_PREFIX + "get_customer_balance_summary", "method": "GET", "desc": "Get customer balance summary"},
    {"path": API_PREFIX + "mark_customer_visit", "method": "POST", "desc": "Mark a customer visit with GPS"},
    {"path": API_PREFIX + "get_customer_visits", "method": "GET", "desc": "Get customer visit history"},
]
add_endpoint_summary_table(cust_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_customer_list", "GET",
    "Get a basic list of customers filtered by name or mobile number.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by customer name"},
        {"name": "mobile_no", "type": "String", "required": "No", "desc": "Filter by mobile number"},
    ],
    response_body={
        "message": [
            {"name": "CUST-0001", "customer_name": "Al-Baik Restaurant", "mobile_no": "+966501234567", "territory": "Riyadh"},
            {"name": "CUST-0002", "customer_name": "Tamimi Markets", "mobile_no": "+966509876543", "territory": "Jeddah"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_list_v2", "GET",
    "Get paginated customer list with advanced filters including territory, route and visit status.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by customer name"},
        {"name": "mobile_no", "type": "String", "required": "No", "desc": "Filter by mobile number"},
        {"name": "territory", "type": "String", "required": "No", "desc": "Filter by territory"},
        {"name": "route", "type": "String", "required": "No", "desc": "Filter by route"},
        {"name": "visit_status", "type": "String", "required": "No", "desc": "Filter by visit status"},
        {"name": "page_number", "type": "Int", "required": "No", "desc": "Page number (default: 1)"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Results per page (default: 20)"},
    ],
    response_body={
        "message": {
            "data": [
                {"name": "CUST-0001", "customer_name": "Al-Baik Restaurant", "mobile_no": "+966501234567", "territory": "Riyadh", "route": "RTE-001", "visit_status": "Visited"},
            ],
            "total_count": 85,
            "page_number": 1,
            "page_size": 20
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_list_with_effective_price_list", "GET",
    "Get customer list with their effective (assigned or default) price list.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by customer name"},
    ],
    response_body={
        "message": [
            {"name": "CUST-0001", "customer_name": "Al-Baik Restaurant", "price_list": "Wholesale Price List"},
            {"name": "CUST-0002", "customer_name": "Tamimi Markets", "price_list": "Standard Selling"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_detail", "POST",
    "Get full details of a specific customer.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID (e.g., CUST-0001)"},
    ],
    request_body={"customer": "CUST-0001"},
    response_body={
        "message": {
            "name": "CUST-0001",
            "customer_name": "Al-Baik Restaurant",
            "customer_group": "Retail",
            "territory": "Riyadh",
            "tax_id": "300012345600003",
            "mobile_no": "+966501234567",
            "email_id": "info@albaik.com",
            "primary_address": "King Fahad Road, Riyadh 12345",
            "outstanding_amount": 12500.00,
            "credit_limit": 50000.00,
            "price_list": "Wholesale Price List"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "create_customer", "POST",
    "Create a new customer with address details.",
    params=[
        {"name": "params", "type": "Object", "required": "Yes", "desc": "Customer details object"},
    ],
    request_body={
        "params": {
            "name": "New Grocery Store",
            "customer_group": "Retail",
            "tax_id": "300098765400003",
            "mobile_no": "+966551234567",
            "email": "newstore@example.com",
            "address": "Al-Olaya Street, Riyadh"
        }
    },
    response_body={
        "message": {
            "name": "CUST-0055",
            "customer_name": "New Grocery Store",
            "customer_group": "Retail"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_invoices", "GET",
    "Get all invoices for a specific customer.",
    params=[
        {"name": "customer", "type": "String", "required": "No", "desc": "Customer ID"},
        {"name": "customer_name", "type": "String", "required": "No", "desc": "Customer name filter"},
    ],
    response_body={
        "message": [
            {"name": "SINV-0001", "customer": "CUST-0001", "grand_total": 15000.00, "outstanding_amount": 5000.00, "status": "Unpaid", "posting_date": "2026-03-15"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "customer_invoice_details", "POST",
    "Get detailed invoice information for a customer.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    request_body={"customer": "CUST-0001"},
    response_body={
        "message": {
            "invoices": [
                {"name": "SINV-0001", "grand_total": 15000.00, "outstanding_amount": 5000.00, "status": "Unpaid", "posting_date": "2026-03-15"}
            ],
            "total_outstanding": 5000.00
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_invoice_aging", "GET",
    "Get invoice aging report across all customers. No parameters required.",
    response_body={
        "message": [
            {"customer": "CUST-0001", "customer_name": "Al-Baik Restaurant", "current": 5000.00, "30_days": 3000.00, "60_days": 0, "90_days": 0, "total_outstanding": 8000.00}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_balance_summary", "GET",
    "Get balance summary for a specific customer.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    response_body={
        "message": {
            "customer": "CUST-0001",
            "customer_name": "Al-Baik Restaurant",
            "total_invoiced": 150000.00,
            "total_paid": 142000.00,
            "outstanding_amount": 8000.00,
            "credit_limit": 50000.00
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "mark_customer_visit", "POST",
    "Record a customer visit with GPS coordinates and notes.",
    params=[
        {"name": "params", "type": "Object", "required": "Yes", "desc": "Visit details object"},
    ],
    request_body={
        "params": {
            "customer": "CUST-0001",
            "notes": "Delivered 50 cases of water. Customer requested more juice next visit.",
            "latitude": 24.7136,
            "longitude": 46.6753
        }
    },
    response_body={
        "message": {
            "name": "VISIT-0001",
            "customer": "CUST-0001",
            "visit_date": "2026-03-21",
            "status": "Completed"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_visits", "GET",
    "Get customer visit history with date range filter.",
    params=[
        {"name": "customer", "type": "String", "required": "No", "desc": "Customer ID"},
        {"name": "from_date", "type": "String", "required": "No", "desc": "Start date (YYYY-MM-DD)"},
        {"name": "to_date", "type": "String", "required": "No", "desc": "End date (YYYY-MM-DD)"},
    ],
    response_body={
        "message": [
            {"name": "VISIT-0001", "customer": "CUST-0001", "visit_date": "2026-03-21", "notes": "Delivered 50 cases.", "latitude": 24.7136, "longitude": 46.6753}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: Customer Route
# ══════════════════════════════════════════════════════════════════════════════

section_heading(4, "Customer Route")

route_endpoints = [
    {"path": API_PREFIX + "get_customer_route_list", "method": "GET", "desc": "Get customer route list with day filters"},
]
add_endpoint_summary_table(route_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_customer_route_list", "GET",
    "Get list of customer routes with day checkboxes and customer count. Can filter for today's routes only.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by route name"},
        {"name": "route_code", "type": "String", "required": "No", "desc": "Filter by route code"},
        {"name": "today_only", "type": "Int", "required": "No", "desc": "Set to 1 to return only today's routes"},
    ],
    response_body={
        "message": [
            {
                "name": "RTE-0001",
                "route_code": "R-RYD-01",
                "route_name": "Riyadh Central Route",
                "sunday": 1,
                "monday": 1,
                "tuesday": 0,
                "wednesday": 1,
                "thursday": 0,
                "friday": 0,
                "saturday": 0,
                "customer_count": 18
            }
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: Items
# ══════════════════════════════════════════════════════════════════════════════

section_heading(5, "Items")

item_endpoints = [
    {"path": API_PREFIX + "get_item_list", "method": "GET", "desc": "Get item list with pagination"},
    {"path": API_PREFIX + "get_item_detail", "method": "POST", "desc": "Get item details"},
    {"path": API_PREFIX + "get_item_by_barcode", "method": "GET", "desc": "Get item by barcode (guest allowed)"},
    {"path": API_PREFIX + "get_available_batches", "method": "GET", "desc": "Get available batches for an item"},
    {"path": API_PREFIX + "sync_items_for_pos", "method": "GET", "desc": "Sync items for POS with delta support"},
    {"path": API_PREFIX + "get_item_stock", "method": "GET", "desc": "Get item stock in a warehouse"},
    {"path": API_PREFIX + "get_item_pricing_details", "method": "GET", "desc": "Get item pricing details"},
    {"path": API_PREFIX + "get_item_uom_details", "method": "GET", "desc": "Get item UOM conversion details"},
]
add_endpoint_summary_table(item_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_item_list", "GET",
    "Get paginated list of items optionally filtered by name, item name, or customer.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by item code"},
        {"name": "item_name", "type": "String", "required": "No", "desc": "Filter by item name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter items by customer (for customer-specific pricing)"},
        {"name": "limit_start", "type": "Int", "required": "No", "desc": "Pagination offset (default: 0)"},
        {"name": "limit_page_length", "type": "Int", "required": "No", "desc": "Results per page (default: 20)"},
    ],
    response_body={
        "message": [
            {"name": "ITEM-0001", "item_name": "Mineral Water 500ml", "item_group": "Beverages", "stock_uom": "Nos", "image": "/files/water500.jpg", "rate": 1.50},
            {"name": "ITEM-0002", "item_name": "Orange Juice 1L", "item_group": "Beverages", "stock_uom": "Nos", "image": "/files/oj1l.jpg", "rate": 5.00}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_item_detail", "POST",
    "Get full details of a specific item.",
    params=[
        {"name": "item", "type": "String", "required": "Yes", "desc": "Item code"},
    ],
    request_body={"item": "ITEM-0001"},
    response_body={
        "message": {
            "name": "ITEM-0001",
            "item_name": "Mineral Water 500ml",
            "item_group": "Beverages",
            "stock_uom": "Nos",
            "description": "Premium mineral water, 500ml bottle",
            "image": "/files/water500.jpg",
            "standard_rate": 1.50,
            "has_batch_no": 1,
            "has_serial_no": 0,
            "barcodes": [{"barcode": "6281234567890"}],
            "uoms": [{"uom": "Nos", "conversion_factor": 1}, {"uom": "Case", "conversion_factor": 24}]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_item_by_barcode", "GET",
    "Lookup an item by its barcode. Allows guest access.",
    params=[
        {"name": "barcode", "type": "String", "required": "Yes", "desc": "Item barcode"},
    ],
    response_body={
        "message": {
            "name": "ITEM-0001",
            "item_name": "Mineral Water 500ml",
            "barcode": "6281234567890",
            "stock_uom": "Nos",
            "rate": 1.50
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_available_batches", "GET",
    "Get available batches for a batch-tracked item in a specific warehouse.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "warehouse", "type": "String", "required": "No", "desc": "Warehouse name"},
    ],
    response_body={
        "message": [
            {"batch_no": "BATCH-2026-001", "expiry_date": "2027-03-01", "qty": 500},
            {"batch_no": "BATCH-2026-002", "expiry_date": "2027-06-15", "qty": 320}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "sync_items_for_pos", "GET",
    "Sync items for POS with delta support. Retrieves items modified since last_sync timestamp.",
    params=[
        {"name": "last_sync", "type": "String", "required": "No", "desc": "Last sync timestamp (ISO format). If empty, returns all items."},
        {"name": "page", "type": "Int", "required": "No", "desc": "Page number (default: 0)"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Items per page (default: 500)"},
    ],
    response_body={
        "message": {
            "items": [
                {"name": "ITEM-0001", "item_name": "Mineral Water 500ml", "rate": 1.50, "stock_uom": "Nos", "modified": "2026-03-20 14:30:00"}
            ],
            "has_more": False,
            "sync_timestamp": "2026-03-21 08:00:00"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_item_stock", "GET",
    "Get current stock quantity for an item in a specific warehouse.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "warehouse", "type": "String", "required": "No", "desc": "Warehouse name"},
    ],
    response_body={
        "message": {
            "item_code": "ITEM-0001",
            "warehouse": "Van Warehouse - AMD",
            "actual_qty": 850,
            "reserved_qty": 50,
            "available_qty": 800
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_item_pricing_details", "GET",
    "Get pricing details for an item, optionally for a specific customer and UOM.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Customer ID for customer-specific pricing"},
        {"name": "uom", "type": "String", "required": "No", "desc": "Unit of measure"},
    ],
    response_body={
        "message": {
            "item_code": "ITEM-0001",
            "price_list": "Wholesale Price List",
            "rate": 1.25,
            "uom": "Nos",
            "currency": "SAR"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_item_uom_details", "GET",
    "Get all UOM conversion details for an item.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
    ],
    response_body={
        "message": [
            {"uom": "Nos", "conversion_factor": 1},
            {"uom": "Case", "conversion_factor": 24},
            {"uom": "Pallet", "conversion_factor": 480}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: Price List
# ══════════════════════════════════════════════════════════════════════════════

section_heading(6, "Price List")

pl_endpoints = [
    {"path": API_PREFIX + "get_price_list", "method": "GET", "desc": "Get list of selling price lists"},
    {"path": API_PREFIX + "get_price_list_details", "method": "GET", "desc": "Get price list details with customers and item count"},
    {"path": API_PREFIX + "get_price_list_items", "method": "GET", "desc": "Get items in a price list"},
    {"path": API_PREFIX + "set_customer_price_list", "method": "POST", "desc": "Assign a price list to a customer"},
    {"path": API_PREFIX + "get_customer_price_list", "method": "GET", "desc": "Get customer's assigned price list"},
]
add_endpoint_summary_table(pl_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_price_list", "GET",
    "Get all available selling price lists. No parameters required.",
    response_body={
        "message": [
            {"name": "Standard Selling", "currency": "SAR"},
            {"name": "Wholesale Price List", "currency": "SAR"},
            {"name": "VIP Customers", "currency": "SAR"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_price_list_details", "GET",
    "Get price list details including assigned customers, item count and companies.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by price list name"},
        {"name": "page_number", "type": "Int", "required": "No", "desc": "Page number (default: 1)"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Results per page (default: 20)"},
    ],
    response_body={
        "message": {
            "data": [
                {
                    "name": "Wholesale Price List",
                    "currency": "SAR",
                    "customers": ["CUST-0001", "CUST-0002"],
                    "item_count": 150,
                    "companies": ["Al-Marai Distribution LLC"]
                }
            ],
            "total_count": 3
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_price_list_items", "GET",
    "Get items in a specific price list with search and pagination.",
    params=[
        {"name": "price_list", "type": "String", "required": "Yes", "desc": "Price list name"},
        {"name": "item_code", "type": "String", "required": "No", "desc": "Filter by item code"},
        {"name": "item_name", "type": "String", "required": "No", "desc": "Filter by item name"},
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by price list item name"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search across name, item_code, item_name"},
        {"name": "page_number", "type": "Int", "required": "No", "desc": "Page number (default: 1)"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Results per page (default: 50)"},
    ],
    response_body={
        "message": {
            "data": [
                {"name": "IPRICE-001", "item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "price_list_rate": 1.25, "currency": "SAR", "uom": "Nos"}
            ],
            "total_count": 150,
            "page_number": 1,
            "page_size": 50
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "set_customer_price_list", "POST",
    "Assign a specific price list to a customer.",
    params=[
        {"name": "params", "type": "Object", "required": "Yes", "desc": "Object with customer and price_list fields"},
    ],
    request_body={
        "params": {
            "customer": "CUST-0001",
            "price_list": "Wholesale Price List"
        }
    },
    response_body={
        "message": "Price list updated successfully"
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_price_list", "GET",
    "Get the assigned price list for a customer.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    response_body={
        "message": {
            "customer": "CUST-0001",
            "price_list": "Wholesale Price List"
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: Sales Order
# ══════════════════════════════════════════════════════════════════════════════

section_heading(7, "Sales Order")

so_endpoints = [
    {"path": API_PREFIX + "create_sales_order", "method": "POST", "desc": "Create a new sales order"},
    {"path": API_PREFIX + "update_sales_order", "method": "POST", "desc": "Update an existing sales order"},
    {"path": API_PREFIX + "submit_sales_order", "method": "POST", "desc": "Submit a draft sales order"},
    {"path": API_PREFIX + "cancel_sales_order", "method": "POST", "desc": "Cancel a submitted sales order"},
    {"path": API_PREFIX + "delete_sales_order", "method": "POST", "desc": "Delete a draft sales order"},
    {"path": API_PREFIX + "get_sales_order_list", "method": "GET", "desc": "Get sales order list with filters"},
    {"path": API_PREFIX + "get_sales_order_detail", "method": "GET", "desc": "Get sales order detail"},
    {"path": API_PREFIX + "convert_so_to_si", "method": "POST", "desc": "Convert sales orders to sales invoice"},
]
add_endpoint_summary_table(so_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_sales_order", "POST",
    "Create a new sales order with items, delivery date and optional price list.",
    params=[
        {"name": "params", "type": "Object", "required": "Yes", "desc": "Sales order details"},
    ],
    request_body={
        "params": {
            "customer": "CUST-0001",
            "delivery_date": "2026-03-25",
            "price_list": "Wholesale Price List",
            "remarks": "Urgent delivery before noon",
            "items": [
                {"item_code": "ITEM-0001", "qty": 100, "uom": "Nos", "rate": 1.25, "discount_percentage": 0, "discount_amount": 0, "is_free_item": 0, "batch_no": "BATCH-2026-001"},
                {"item_code": "ITEM-0002", "qty": 5, "uom": "Case", "rate": 110.00, "discount_percentage": 5, "discount_amount": 0, "is_free_item": 0, "batch_no": ""}
            ]
        }
    },
    response_body={
        "message": {
            "name": "SO-0045",
            "customer": "CUST-0001",
            "grand_total": 647.50,
            "status": "Draft",
            "delivery_date": "2026-03-25"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "update_sales_order", "POST",
    "Update an existing draft sales order.",
    params=[
        {"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name, items and delivery_date"},
    ],
    request_body={
        "params": {
            "name": "SO-0045",
            "delivery_date": "2026-03-26",
            "items": [
                {"item_code": "ITEM-0001", "qty": 120, "uom": "Nos", "rate": 1.25}
            ]
        }
    },
    response_body={"message": {"name": "SO-0045", "status": "Draft", "grand_total": 150.00}},
)

add_endpoint_detail(
    API_PREFIX + "submit_sales_order", "POST",
    "Submit a draft sales order for processing.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SO-0045"}},
    response_body={"message": {"name": "SO-0045", "status": "To Deliver and Bill", "docstatus": 1}},
)

add_endpoint_detail(
    API_PREFIX + "cancel_sales_order", "POST",
    "Cancel a submitted sales order.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SO-0045"}},
    response_body={"message": {"name": "SO-0045", "status": "Cancelled", "docstatus": 2}},
)

add_endpoint_detail(
    API_PREFIX + "delete_sales_order", "POST",
    "Permanently delete a draft sales order.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SO-0045"}},
    response_body={"message": "Sales Order SO-0045 deleted successfully"},
)

add_endpoint_detail(
    API_PREFIX + "get_sales_order_list", "GET",
    "Get paginated list of sales orders with filters.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by SO name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer ID"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status (Draft, Submitted, Cancelled)"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search"},
        {"name": "from_date", "type": "String", "required": "No", "desc": "Start date filter (YYYY-MM-DD)"},
        {"name": "to_date", "type": "String", "required": "No", "desc": "End date filter (YYYY-MM-DD)"},
        {"name": "route", "type": "String", "required": "No", "desc": "Filter by route"},
        {"name": "page_number", "type": "Int", "required": "No", "desc": "Page number (default: 1)"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Results per page (default: 20)"},
    ],
    response_body={
        "message": {
            "data": [
                {
                    "name": "SO-0045",
                    "customer": "CUST-0001",
                    "customer_name": "Al-Baik Restaurant",
                    "customer_address": "King Fahad Road, Riyadh",
                    "custom_route": "RTE-001",
                    "grand_total": 647.50,
                    "status": "To Deliver and Bill",
                    "created_date": "2026-03-21",
                    "created_by": "salesman@example.com",
                    "items": [
                        {"item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "qty": 100, "rate": 1.25, "amount": 125.00, "discount_percentage": 0, "discount_amount": 0, "conversion_factor": 1}
                    ]
                }
            ],
            "total_count": 45
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_sales_order_detail", "GET",
    "Get full details of a specific sales order.",
    params=[{"name": "name", "type": "String", "required": "Yes", "desc": "Sales order name"}],
    response_body={
        "message": {
            "name": "SO-0045",
            "customer": "CUST-0001",
            "customer_name": "Al-Baik Restaurant",
            "delivery_date": "2026-03-25",
            "grand_total": 647.50,
            "status": "To Deliver and Bill",
            "remarks": "Urgent delivery before noon",
            "items": [
                {"item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "qty": 100, "uom": "Nos", "rate": 1.25, "amount": 125.00}
            ]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "convert_so_to_si", "POST",
    "Convert one or more submitted sales orders to a sales invoice, with optional payment details.",
    params=[
        {"name": "params", "type": "Object", "required": "Yes", "desc": "Object with sales_orders list and optional payments"},
    ],
    request_body={
        "params": {
            "sales_orders": ["SO-0045", "SO-0046"],
            "payments": [
                {"mode_of_payment": "Cash", "amount": 500.00}
            ]
        }
    },
    response_body={
        "message": {
            "sales_invoice": "SINV-0055",
            "grand_total": 1200.00,
            "status": "Draft"
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: Delivery Note (Optional)
# ══════════════════════════════════════════════════════════════════════════════

section_heading(8, "Delivery Note (Optional)")

doc.add_paragraph(
    "This section is Optional and only applies when 'Ignore Create Delivery Note' is unchecked "
    "in Basket4Me Settings. When enabled, delivery notes are used as an intermediate step "
    "between Sales Order and Sales Invoice."
)
p_note = doc.add_paragraph()
run_note = p_note.add_run("Note: ")
run_note.bold = True
p_note.add_run("If 'Ignore Create Delivery Note' is checked, skip this section entirely and convert SO directly to SI.")

dn_endpoints = [
    {"path": API_PREFIX + "create_delivery_note", "method": "POST", "desc": "Create a new delivery note"},
    {"path": API_PREFIX + "update_delivery_note", "method": "POST", "desc": "Update an existing delivery note"},
    {"path": API_PREFIX + "submit_delivery_note", "method": "POST", "desc": "Submit a delivery note"},
    {"path": API_PREFIX + "cancel_delivery_note", "method": "POST", "desc": "Cancel a delivery note"},
    {"path": API_PREFIX + "delete_delivery_note", "method": "POST", "desc": "Delete a draft delivery note"},
    {"path": API_PREFIX + "get_delivery_note_list", "method": "GET", "desc": "Get delivery note list with filters"},
    {"path": API_PREFIX + "get_delivery_note_detail", "method": "GET", "desc": "Get delivery note detail"},
    {"path": API_PREFIX + "get_pending_so_for_dn", "method": "GET", "desc": "Get pending SOs for delivery note creation"},
    {"path": API_PREFIX + "create_dn_from_so", "method": "POST", "desc": "Create delivery note from sales order(s)"},
    {"path": API_PREFIX + "get_uninvoiced_dn", "method": "GET", "desc": "Get uninvoiced delivery notes for a customer"},
]
add_endpoint_summary_table(dn_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_delivery_note", "POST",
    "Create a new delivery note with customer and items.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Delivery note details with customer and items"}],
    request_body={
        "params": {
            "customer": "CUST-0001",
            "items": [
                {"item_code": "ITEM-0001", "qty": 100, "uom": "Nos", "rate": 1.25, "warehouse": "Van Warehouse - AMD"}
            ]
        }
    },
    response_body={"message": {"name": "DN-0012", "customer": "CUST-0001", "grand_total": 125.00, "status": "Draft"}},
)

add_endpoint_detail(
    API_PREFIX + "update_delivery_note", "POST",
    "Update an existing draft delivery note.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name and items"}],
    request_body={
        "params": {
            "name": "DN-0012",
            "items": [{"item_code": "ITEM-0001", "qty": 120, "uom": "Nos", "rate": 1.25}]
        }
    },
    response_body={"message": {"name": "DN-0012", "status": "Draft", "grand_total": 150.00}},
)

add_endpoint_detail(
    API_PREFIX + "submit_delivery_note", "POST",
    "Submit a draft delivery note.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "DN-0012"}},
    response_body={"message": {"name": "DN-0012", "status": "To Bill", "docstatus": 1}},
)

add_endpoint_detail(
    API_PREFIX + "cancel_delivery_note", "POST",
    "Cancel a submitted delivery note.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "DN-0012"}},
    response_body={"message": {"name": "DN-0012", "status": "Cancelled", "docstatus": 2}},
)

add_endpoint_detail(
    API_PREFIX + "delete_delivery_note", "POST",
    "Delete a draft delivery note.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "DN-0012"}},
    response_body={"message": "Delivery Note DN-0012 deleted successfully"},
)

add_endpoint_detail(
    API_PREFIX + "get_delivery_note_list", "GET",
    "Get paginated list of delivery notes with filters.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by DN name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search"},
        {"name": "from_date", "type": "String", "required": "No", "desc": "Start date (YYYY-MM-DD)"},
        {"name": "to_date", "type": "String", "required": "No", "desc": "End date (YYYY-MM-DD)"},
        {"name": "page_number", "type": "Int", "required": "No", "desc": "Page number"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Results per page"},
    ],
    response_body={
        "message": {
            "data": [
                {"name": "DN-0012", "customer": "CUST-0001", "customer_name": "Al-Baik Restaurant", "grand_total": 125.00, "status": "To Bill", "posting_date": "2026-03-21"}
            ],
            "total_count": 22
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_delivery_note_detail", "GET",
    "Get full details of a specific delivery note.",
    params=[{"name": "name", "type": "String", "required": "Yes", "desc": "Delivery note name"}],
    response_body={
        "message": {
            "name": "DN-0012",
            "customer": "CUST-0001",
            "customer_name": "Al-Baik Restaurant",
            "grand_total": 125.00,
            "status": "To Bill",
            "items": [{"item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "qty": 100, "rate": 1.25, "amount": 125.00}]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_pending_so_for_dn", "GET",
    "Get sales orders pending delivery for a specific customer.",
    params=[{"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"}],
    response_body={
        "message": [
            {"name": "SO-0045", "customer": "CUST-0001", "grand_total": 647.50, "delivery_date": "2026-03-25", "per_delivered": 0}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "create_dn_from_so", "POST",
    "Create a delivery note from one or more sales orders.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with sales_order or sales_orders list"}],
    request_body={"params": {"sales_orders": ["SO-0045"]}},
    response_body={"message": {"name": "DN-0013", "customer": "CUST-0001", "grand_total": 647.50, "status": "Draft"}},
)

add_endpoint_detail(
    API_PREFIX + "get_uninvoiced_dn", "GET",
    "Get delivery notes that have not yet been invoiced for a customer.",
    params=[{"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"}],
    response_body={
        "message": [
            {"name": "DN-0012", "customer": "CUST-0001", "grand_total": 125.00, "status": "To Bill", "posting_date": "2026-03-21"}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9: Sales Invoice
# ══════════════════════════════════════════════════════════════════════════════

section_heading(9, "Sales Invoice")

si_endpoints = [
    {"path": API_PREFIX + "create_sales_invoice", "method": "POST", "desc": "Create a new sales invoice"},
    {"path": API_PREFIX + "update_sales_invoice", "method": "POST", "desc": "Update an existing sales invoice"},
    {"path": API_PREFIX + "submit_sales_invoice", "method": "POST", "desc": "Submit a draft sales invoice"},
    {"path": API_PREFIX + "cancel_sales_invoice", "method": "POST", "desc": "Cancel a submitted sales invoice"},
    {"path": API_PREFIX + "delete_sales_invoice", "method": "POST", "desc": "Delete a draft sales invoice"},
    {"path": API_PREFIX + "get_invoice_list", "method": "GET", "desc": "Get sales invoice list with filters"},
    {"path": API_PREFIX + "get_invoice_detail", "method": "GET", "desc": "Get sales invoice detail"},
    {"path": API_PREFIX + "invoice_details", "method": "POST", "desc": "Get invoice details by ID"},
    {"path": API_PREFIX + "delete_invoice", "method": "DELETE", "desc": "Delete an invoice (alternative)"},
    {"path": API_PREFIX + "create_si_from_dn", "method": "POST", "desc": "Create sales invoice from delivery notes"},
]
add_endpoint_summary_table(si_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_sales_invoice", "POST",
    "Create a new sales invoice with items and optional price list.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Invoice details with customer, posting_date, due_date, items, price_list"}],
    request_body={
        "params": {
            "customer": "CUST-0001",
            "posting_date": "2026-03-21",
            "due_date": "2026-04-20",
            "price_list": "Wholesale Price List",
            "items": [
                {"item_code": "ITEM-0001", "qty": 100, "uom": "Nos", "rate": 1.25},
                {"item_code": "ITEM-0002", "qty": 5, "uom": "Case", "rate": 110.00}
            ]
        }
    },
    response_body={
        "message": {
            "name": "SINV-0055",
            "customer": "CUST-0001",
            "grand_total": 675.00,
            "status": "Draft",
            "posting_date": "2026-03-21"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "update_sales_invoice", "POST",
    "Update a draft sales invoice.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name and items"}],
    request_body={
        "params": {
            "name": "SINV-0055",
            "items": [{"item_code": "ITEM-0001", "qty": 150, "uom": "Nos", "rate": 1.25}]
        }
    },
    response_body={"message": {"name": "SINV-0055", "status": "Draft", "grand_total": 187.50}},
)

add_endpoint_detail(
    API_PREFIX + "submit_sales_invoice", "POST",
    "Submit a draft sales invoice for accounting.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SINV-0055"}},
    response_body={"message": {"name": "SINV-0055", "status": "Unpaid", "docstatus": 1}},
)

add_endpoint_detail(
    API_PREFIX + "cancel_sales_invoice", "POST",
    "Cancel a submitted sales invoice.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SINV-0055"}},
    response_body={"message": {"name": "SINV-0055", "status": "Cancelled", "docstatus": 2}},
)

add_endpoint_detail(
    API_PREFIX + "delete_sales_invoice", "POST",
    "Delete a draft sales invoice.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SINV-0055"}},
    response_body={"message": "Sales Invoice SINV-0055 deleted successfully"},
)

add_endpoint_detail(
    API_PREFIX + "get_invoice_list", "GET",
    "Get list of sales invoices with filters.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by invoice name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search"},
    ],
    response_body={
        "message": [
            {"name": "SINV-0055", "customer": "CUST-0001", "customer_name": "Al-Baik Restaurant", "grand_total": 675.00, "outstanding_amount": 675.00, "status": "Unpaid", "posting_date": "2026-03-21"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_invoice_detail", "GET",
    "Get full details of a specific sales invoice.",
    params=[{"name": "name", "type": "String", "required": "Yes", "desc": "Invoice name"}],
    response_body={
        "message": {
            "name": "SINV-0055",
            "customer": "CUST-0001",
            "customer_name": "Al-Baik Restaurant",
            "posting_date": "2026-03-21",
            "due_date": "2026-04-20",
            "grand_total": 675.00,
            "outstanding_amount": 675.00,
            "status": "Unpaid",
            "items": [
                {"item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "qty": 100, "rate": 1.25, "amount": 125.00},
                {"item_code": "ITEM-0002", "item_name": "Orange Juice 1L", "qty": 5, "rate": 110.00, "amount": 550.00}
            ]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "invoice_details", "POST",
    "Get invoice details by invoice ID (alternative endpoint).",
    params=[{"name": "invoice_id", "type": "String", "required": "Yes", "desc": "Invoice name/ID"}],
    request_body={"invoice_id": "SINV-0055"},
    response_body={
        "message": {
            "name": "SINV-0055",
            "customer": "CUST-0001",
            "grand_total": 675.00,
            "outstanding_amount": 675.00,
            "status": "Unpaid"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "delete_invoice", "DELETE",
    "Delete an invoice using DELETE method (alternative endpoint).",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "SINV-0055"}},
    response_body={"message": "Invoice SINV-0055 deleted successfully"},
)

add_endpoint_detail(
    API_PREFIX + "create_si_from_dn", "POST",
    "Create a sales invoice from one or more delivery notes.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with delivery_notes list"}],
    request_body={"params": {"delivery_notes": ["DN-0012", "DN-0013"]}},
    response_body={
        "message": {
            "sales_invoice": "SINV-0060",
            "grand_total": 772.50,
            "status": "Draft"
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10: Sales Return
# ══════════════════════════════════════════════════════════════════════════════

section_heading(10, "Sales Return")

ret_endpoints = [
    {"path": API_PREFIX + "create_sales_invoice_return", "method": "POST", "desc": "Create a sales return (credit note)"},
    {"path": API_PREFIX + "get_return_invoice_list", "method": "GET", "desc": "Get return invoice list"},
    {"path": API_PREFIX + "get_returned_qty", "method": "GET", "desc": "Get returned quantities for an invoice"},
    {"path": API_PREFIX + "get_invoices_for_return", "method": "GET", "desc": "Get invoices eligible for return"},
]
add_endpoint_summary_table(ret_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_sales_invoice_return", "POST",
    "Create a sales return (credit note) against a submitted invoice. Quantities must be negative.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Return details with return_against and items"}],
    request_body={
        "params": {
            "return_against": "SINV-0055",
            "items": [
                {"item_code": "ITEM-0001", "qty": -20, "uom": "Nos"},
                {"item_code": "ITEM-0002", "qty": -1, "uom": "Case"}
            ]
        }
    },
    response_body={
        "message": {
            "name": "SINV-RET-0001",
            "return_against": "SINV-0055",
            "grand_total": -135.00,
            "status": "Return"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_return_invoice_list", "GET",
    "Get list of return invoices (credit notes).",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by return name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search"},
    ],
    response_body={
        "message": [
            {"name": "SINV-RET-0001", "customer": "CUST-0001", "return_against": "SINV-0055", "grand_total": -135.00, "status": "Return"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_returned_qty", "GET",
    "Get the quantities already returned for each item in a specific invoice.",
    params=[{"name": "invoice_name", "type": "String", "required": "Yes", "desc": "Original invoice name"}],
    response_body={
        "message": [
            {"item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "invoiced_qty": 100, "returned_qty": 20, "returnable_qty": 80},
            {"item_code": "ITEM-0002", "item_name": "Orange Juice 1L", "invoiced_qty": 5, "returned_qty": 1, "returnable_qty": 4}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_invoices_for_return", "GET",
    "Get submitted invoices eligible for return for a specific customer.",
    params=[{"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"}],
    response_body={
        "message": [
            {"name": "SINV-0055", "customer": "CUST-0001", "grand_total": 675.00, "outstanding_amount": 675.00, "posting_date": "2026-03-21"}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11: Receipts / Payments
# ══════════════════════════════════════════════════════════════════════════════

section_heading(11, "Receipts / Payments")

pay_endpoints = [
    {"path": API_PREFIX + "create_payment_entry", "method": "POST", "desc": "Create a payment entry"},
    {"path": API_PREFIX + "cancel_payment_entry", "method": "POST", "desc": "Cancel a payment entry"},
    {"path": API_PREFIX + "get_receipt_list", "method": "GET", "desc": "Get receipt/payment list"},
    {"path": API_PREFIX + "receipt_details", "method": "POST", "desc": "Get receipt details"},
    {"path": API_PREFIX + "get_available_modes_of_payment", "method": "GET", "desc": "Get available modes of payment"},
]
add_endpoint_summary_table(pay_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_payment_entry", "POST",
    "Create a payment entry with references to invoices.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Payment details with customer, amount, mode and references"}],
    request_body={
        "params": {
            "customer": "CUST-0001",
            "paid_amount": 500.00,
            "mode_of_payment": "Cash",
            "reference_no": "REC-2026-0001",
            "references": [
                {"reference_doctype": "Sales Invoice", "reference_name": "SINV-0055", "allocated_amount": 500.00}
            ]
        }
    },
    response_body={
        "message": {
            "name": "PE-0032",
            "customer": "CUST-0001",
            "paid_amount": 500.00,
            "mode_of_payment": "Cash",
            "status": "Submitted"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "cancel_payment_entry", "POST",
    "Cancel a submitted payment entry.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "PE-0032"}},
    response_body={"message": {"name": "PE-0032", "status": "Cancelled", "docstatus": 2}},
)

add_endpoint_detail(
    API_PREFIX + "get_receipt_list", "GET",
    "Get paginated list of receipts/payments.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by receipt name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search"},
        {"name": "page", "type": "Int", "required": "No", "desc": "Page number (default: 1)"},
        {"name": "page_size", "type": "Int", "required": "No", "desc": "Results per page (default: 20)"},
    ],
    response_body={
        "message": [
            {"name": "PE-0032", "customer": "CUST-0001", "paid_amount": 500.00, "mode_of_payment": "Cash", "status": "Submitted", "posting_date": "2026-03-21"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "receipt_details", "POST",
    "Get full details of a specific receipt/payment.",
    params=[{"name": "receipt_id", "type": "String", "required": "Yes", "desc": "Receipt/Payment entry name"}],
    request_body={"receipt_id": "PE-0032"},
    response_body={
        "message": {
            "name": "PE-0032",
            "customer": "CUST-0001",
            "paid_amount": 500.00,
            "mode_of_payment": "Cash",
            "reference_no": "REC-2026-0001",
            "references": [
                {"reference_doctype": "Sales Invoice", "reference_name": "SINV-0055", "allocated_amount": 500.00}
            ]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_available_modes_of_payment", "GET",
    "Get all available modes of payment. No parameters required.",
    response_body={
        "message": [
            {"name": "Cash"},
            {"name": "Bank Transfer"},
            {"name": "Credit Card"},
            {"name": "Cheque"}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12: Material Request
# ══════════════════════════════════════════════════════════════════════════════

section_heading(12, "Material Request")

mr_endpoints = [
    {"path": API_PREFIX + "create_material_request", "method": "POST", "desc": "Create a material request"},
    {"path": API_PREFIX + "get_material_request_list", "method": "GET", "desc": "Get material request list"},
    {"path": API_PREFIX + "get_material_request_detail", "method": "GET", "desc": "Get material request detail"},
    {"path": API_PREFIX + "update_material_request", "method": "POST", "desc": "Update a material request"},
    {"path": API_PREFIX + "submit_material_request", "method": "POST", "desc": "Submit a material request"},
    {"path": API_PREFIX + "delete_material_request", "method": "POST", "desc": "Delete a material request"},
]
add_endpoint_summary_table(mr_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_material_request", "POST",
    "Create a material request (e.g., van stock replenishment).",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with items list"}],
    request_body={
        "params": {
            "items": [
                {"item_code": "ITEM-0001", "qty": 500, "uom": "Nos", "warehouse": "Van Warehouse - AMD"},
                {"item_code": "ITEM-0002", "qty": 50, "uom": "Case", "warehouse": "Van Warehouse - AMD"}
            ]
        }
    },
    response_body={
        "message": {
            "name": "MR-0010",
            "status": "Draft",
            "items": [
                {"item_code": "ITEM-0001", "qty": 500},
                {"item_code": "ITEM-0002", "qty": 50}
            ]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_material_request_list", "GET",
    "Get list of material requests.",
    params=[{"name": "name", "type": "String", "required": "No", "desc": "Filter by material request name"}],
    response_body={
        "message": [
            {"name": "MR-0010", "status": "Pending", "transaction_date": "2026-03-21", "material_request_type": "Material Transfer"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_material_request_detail", "GET",
    "Get full details of a specific material request.",
    params=[{"name": "name", "type": "String", "required": "Yes", "desc": "Material request name"}],
    response_body={
        "message": {
            "name": "MR-0010",
            "status": "Pending",
            "transaction_date": "2026-03-21",
            "items": [
                {"item_code": "ITEM-0001", "item_name": "Mineral Water 500ml", "qty": 500, "uom": "Nos", "warehouse": "Van Warehouse - AMD"}
            ]
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "update_material_request", "POST",
    "Update an existing material request.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name and updated fields"}],
    request_body={"params": {"name": "MR-0010", "items": [{"item_code": "ITEM-0001", "qty": 600}]}},
    response_body={"message": {"name": "MR-0010", "status": "Draft"}},
)

add_endpoint_detail(
    API_PREFIX + "submit_material_request", "POST",
    "Submit a material request for processing.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "MR-0010"}},
    response_body={"message": {"name": "MR-0010", "status": "Pending", "docstatus": 1}},
)

add_endpoint_detail(
    API_PREFIX + "delete_material_request", "POST",
    "Delete a draft material request.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with name field"}],
    request_body={"params": {"name": "MR-0010"}},
    response_body={"message": "Material Request MR-0010 deleted successfully"},
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 13: Print
# ══════════════════════════════════════════════════════════════════════════════

section_heading(13, "Print")

print_endpoints = [
    {"path": API_PREFIX + "get_print_pdf", "method": "GET/POST", "desc": "Get a print PDF for a document"},
    {"path": API_PREFIX + "get_print_formats", "method": "GET", "desc": "Get available print formats for a DocType"},
    {"path": API_PREFIX + "get_available_print_formats", "method": "GET", "desc": "Get available print formats (alternative)"},
    {"path": API_PREFIX + "print_document", "method": "GET/POST", "desc": "Print/render a document"},
    {"path": API_PREFIX + "download_invoice_pdf", "method": "GET", "desc": "Download invoice PDF directly"},
]
add_endpoint_summary_table(print_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_print_pdf", "GET/POST",
    "Generate and return a PDF for the specified document.",
    params=[
        {"name": "doctype", "type": "String", "required": "Yes", "desc": "DocType (e.g., Sales Invoice)"},
        {"name": "name", "type": "String", "required": "Yes", "desc": "Document name"},
        {"name": "print_format", "type": "String", "required": "No", "desc": "Print format name"},
        {"name": "letterhead", "type": "String", "required": "No", "desc": "Letterhead name"},
    ],
    response_body={"message": "Binary PDF content (Content-Type: application/pdf)"},
)

add_endpoint_detail(
    API_PREFIX + "get_print_formats", "GET",
    "Get all print formats available for a specific DocType.",
    params=[{"name": "doctype", "type": "String", "required": "Yes", "desc": "DocType name"}],
    response_body={
        "message": [
            {"name": "Standard", "doc_type": "Sales Invoice"},
            {"name": "POS Invoice", "doc_type": "Sales Invoice"},
            {"name": "Tax Invoice (KSA)", "doc_type": "Sales Invoice"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_available_print_formats", "GET",
    "Get available print formats for a DocType (alternative endpoint).",
    params=[{"name": "doctype", "type": "String", "required": "Yes", "desc": "DocType name"}],
    response_body={
        "message": ["Standard", "POS Invoice", "Tax Invoice (KSA)"]
    },
)

add_endpoint_detail(
    API_PREFIX + "print_document", "GET/POST",
    "Render and optionally print a document in various output types.",
    params=[
        {"name": "doctype", "type": "String", "required": "Yes", "desc": "DocType name"},
        {"name": "name", "type": "String", "required": "Yes", "desc": "Document name"},
        {"name": "print_format", "type": "String", "required": "No", "desc": "Print format name"},
        {"name": "output_type", "type": "String", "required": "No", "desc": "Output type: 'pdf' or 'html'"},
    ],
    response_body={"message": "PDF binary or HTML string based on output_type"},
)

add_endpoint_detail(
    API_PREFIX + "download_invoice_pdf", "GET",
    "Download a sales invoice PDF directly.",
    params=[
        {"name": "invoice_name", "type": "String", "required": "Yes", "desc": "Sales invoice name"},
        {"name": "print_format", "type": "String", "required": "No", "desc": "Print format to use"},
    ],
    response_body={"message": "Binary PDF file download"},
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 14: Config & Settings
# ══════════════════════════════════════════════════════════════════════════════

section_heading(14, "Config & Settings")

config_endpoints = [
    {"path": API_PREFIX + "get_salesperson_config", "method": "GET", "desc": "Get salesperson configuration"},
    {"path": API_PREFIX + "get_payment_type", "method": "GET", "desc": "Get payment type options"},
    {"path": API_PREFIX + "get_warehouse_list", "method": "GET", "desc": "Get warehouse list"},
    {"path": API_PREFIX + "get_customer_group_list", "method": "GET", "desc": "Get customer group list"},
    {"path": API_PREFIX + "get_territories", "method": "GET", "desc": "Get territory list"},
    {"path": API_PREFIX + "enable_customer_override", "method": "POST", "desc": "Enable customer override setting"},
    {"path": API_PREFIX + "enable_customer_based_price_list", "method": "POST", "desc": "Enable customer-based price list"},
    {"path": API_PREFIX + "toggle_sales_team_override", "method": "POST", "desc": "Toggle sales team override"},
    {"path": API_PREFIX + "get_sales_team_override_status", "method": "GET", "desc": "Get sales team override status"},
    {"path": API_PREFIX + "check_basket4me_settings_status", "method": "GET", "desc": "Check Basket4Me settings status"},
]
add_endpoint_summary_table(config_endpoints)

add_endpoint_detail(
    API_PREFIX + "get_salesperson_config", "GET",
    "Get the salesperson configuration for the current user. No parameters required.",
    response_body={
        "message": {
            "salesperson": "Ahmed Al-Saleh",
            "warehouse": "Van Warehouse - AMD",
            "company": "Al-Marai Distribution LLC",
            "territory": "Riyadh",
            "default_price_list": "Wholesale Price List",
            "ignore_create_delivery_note": 1,
            "enable_customer_override": 0,
            "enable_global_price_list": 1
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_payment_type", "GET",
    "Get available payment type options. No parameters required.",
    response_body={
        "message": [
            {"name": "Cash", "type": "Cash"},
            {"name": "Bank Transfer", "type": "Bank"},
            {"name": "Credit Card", "type": "Bank"},
            {"name": "Cheque", "type": "Bank"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_warehouse_list", "GET",
    "Get list of available warehouses.",
    params=[{"name": "name", "type": "String", "required": "No", "desc": "Filter by warehouse name"}],
    response_body={
        "message": [
            {"name": "Van Warehouse - AMD", "company": "Al-Marai Distribution LLC"},
            {"name": "Main Store - AMD", "company": "Al-Marai Distribution LLC"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_customer_group_list", "GET",
    "Get list of customer groups.",
    params=[{"name": "name", "type": "String", "required": "No", "desc": "Filter by customer group name"}],
    response_body={
        "message": [
            {"name": "Retail"},
            {"name": "Wholesale"},
            {"name": "Distributor"},
            {"name": "Key Account"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "get_territories", "GET",
    "Get list of territories. No parameters required.",
    response_body={
        "message": [
            {"name": "Riyadh"},
            {"name": "Jeddah"},
            {"name": "Dammam"},
            {"name": "Makkah"}
        ]
    },
)

add_endpoint_detail(
    API_PREFIX + "enable_customer_override", "POST",
    "Enable or disable customer override setting.",
    request_body={"params": {"enable": 1}},
    response_body={"message": "Customer override enabled"},
)

add_endpoint_detail(
    API_PREFIX + "enable_customer_based_price_list", "POST",
    "Enable or disable customer-based price list functionality.",
    request_body={"params": {"enable": 1}},
    response_body={"message": "Customer-based price list enabled"},
)

add_endpoint_detail(
    API_PREFIX + "toggle_sales_team_override", "POST",
    "Toggle sales team override on or off.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Toggle parameters"}],
    request_body={"params": {"enable": 1}},
    response_body={"message": "Sales team override enabled"},
)

add_endpoint_detail(
    API_PREFIX + "get_sales_team_override_status", "GET",
    "Get current sales team override status. No parameters required.",
    response_body={
        "message": {
            "enabled": 1,
            "override_by": "salesman@example.com"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "check_basket4me_settings_status", "GET",
    "Check the current Basket4Me settings status. No parameters required.",
    response_body={
        "message": {
            "ignore_create_delivery_note": 1,
            "enable_customer_override": 0,
            "enable_global_price_list": 1,
            "enable_customer_based_price_list": 0,
            "default_warehouse": "Van Warehouse - AMD",
            "default_price_list": "Wholesale Price List"
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 15: Shift Management
# ══════════════════════════════════════════════════════════════════════════════

section_heading(15, "Shift Management")

shift_endpoints = [
    {"path": API_PREFIX + "create_shift_opening_entry", "method": "POST", "desc": "Create a shift opening entry"},
    {"path": API_PREFIX + "create_shift_closing_entry", "method": "POST", "desc": "Create a shift closing entry"},
    {"path": API_PREFIX + "get_open_shift", "method": "GET", "desc": "Get current open shift"},
]
add_endpoint_summary_table(shift_endpoints)

add_endpoint_detail(
    API_PREFIX + "create_shift_opening_entry", "POST",
    "Create a shift opening entry with the opening cash amount.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with opening_amount"}],
    request_body={"params": {"opening_amount": 500.00}},
    response_body={
        "message": {
            "name": "SHIFT-OPEN-0001",
            "opening_amount": 500.00,
            "opening_time": "2026-03-21 08:00:00",
            "status": "Open"
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "create_shift_closing_entry", "POST",
    "Close the current shift with the closing cash amount.",
    params=[{"name": "params", "type": "Object", "required": "Yes", "desc": "Object with closing_amount"}],
    request_body={"params": {"closing_amount": 3500.00}},
    response_body={
        "message": {
            "name": "SHIFT-CLOSE-0001",
            "closing_amount": 3500.00,
            "closing_time": "2026-03-21 18:00:00",
            "status": "Closed",
            "total_sales": 3000.00,
            "difference": 0.00
        }
    },
)

add_endpoint_detail(
    API_PREFIX + "get_open_shift", "GET",
    "Get the currently open shift for the logged-in user. No parameters required.",
    response_body={
        "message": {
            "name": "SHIFT-OPEN-0001",
            "opening_amount": 500.00,
            "opening_time": "2026-03-21 08:00:00",
            "status": "Open",
            "total_sales_so_far": 2100.00,
            "total_collections_so_far": 1800.00
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A: Auth Header Format
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix A: Auth Header Format", level=1)

doc.add_paragraph(
    "All authenticated API requests must include the Authorization header. "
    "The token is obtained from the login endpoint response."
)

doc.add_heading("Header Format", level=2)
make_table(
    ["Header", "Value"],
    [
        ["Authorization", "token <api_key>:<api_secret>"],
        ["Content-Type", "application/json"],
    ],
)

doc.add_paragraph()
doc.add_heading("Example (cURL)", level=2)
p_curl = doc.add_paragraph()
curl_text = (
    'curl -X GET \\\n'
    f'  "{BASE_URL}/api/method/basket4me_pwa.api.get_customer_list" \\\n'
    '  -H "Authorization: token c12a33bb670b1e1:80c7720caff3272" \\\n'
    '  -H "Content-Type: application/json"'
)
run_curl = p_curl.add_run(curl_text)
run_curl.font.name = "Consolas"
run_curl.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("Example (JavaScript / frappe.call)", level=2)
p_js = doc.add_paragraph()
js_text = (
    'frappe.call({\n'
    '  method: "basket4me_pwa.api.get_customer_list",\n'
    '  args: { name: "Al-Baik" },\n'
    '  callback: function(r) {\n'
    '    console.log(r.message);\n'
    '  }\n'
    '});'
)
run_js = p_js.add_run(js_text)
run_js.font.name = "Consolas"
run_js.font.size = Pt(9)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B: Error Response Format
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix B: Standard Error Response Format", level=1)

doc.add_paragraph(
    "All API errors follow the standard Frappe error response format. "
    "The HTTP status code indicates the general category of error."
)

add_json_block("Error Response Example:", {
    "exc_type": "ValidationError",
    "exc": "traceback string (only in development mode)",
    "_server_messages": '[{"message": "Customer CUST-9999 not found", "indicator": "red"}]',
    "message": "Customer CUST-9999 not found"
})

doc.add_paragraph()
doc.add_heading("Common Error Types", level=2)
make_table(
    ["exc_type", "HTTP Status", "Description"],
    [
        ["ValidationError", "417", "Input validation failed"],
        ["PermissionError", "403", "User lacks required permissions"],
        ["DoesNotExistError", "404", "Requested resource not found"],
        ["AuthenticationError", "401", "Invalid or expired credentials"],
        ["DuplicateEntryError", "409", "Record already exists"],
        ["MandatoryError", "417", "Required field is missing"],
        ["TimestampMismatchError", "409", "Document was modified by another user"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX C: HTTP Status Codes
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix C: HTTP Status Codes", level=1)

make_table(
    ["Code", "Status", "Description"],
    [
        ["200", "OK", "Request succeeded"],
        ["201", "Created", "Resource created successfully"],
        ["400", "Bad Request", "Malformed request or missing parameters"],
        ["401", "Unauthorized", "Invalid or missing authentication credentials"],
        ["403", "Forbidden", "User does not have required permissions"],
        ["404", "Not Found", "Requested resource does not exist"],
        ["409", "Conflict", "Duplicate entry or timestamp mismatch"],
        ["417", "Expectation Failed", "Validation error (Frappe standard)"],
        ["500", "Internal Server Error", "Unexpected server-side error"],
        ["502", "Bad Gateway", "Upstream server error (proxy/load balancer)"],
        ["503", "Service Unavailable", "Server temporarily unavailable"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX D: Workflow Diagrams
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix D: Workflow Diagrams", level=1)

doc.add_heading("Workflow 1: Sales Order to Invoice to Receipt", level=2)
doc.add_paragraph(
    "This is the standard van sales workflow when 'Ignore Create Delivery Note' is enabled:"
)
flow1 = doc.add_paragraph()
r1 = flow1.add_run(
    "Sales Order (Draft)\n"
    "    --> submit_sales_order\n"
    "Sales Order (Submitted)\n"
    "    --> convert_so_to_si\n"
    "Sales Invoice (Draft)\n"
    "    --> submit_sales_invoice\n"
    "Sales Invoice (Unpaid)\n"
    "    --> create_payment_entry\n"
    "Payment Entry (Submitted) + Sales Invoice (Paid)"
)
r1.font.name = "Consolas"
r1.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("Workflow 2: Sales Invoice to Return to Receipt", level=2)
doc.add_paragraph(
    "This workflow handles product returns and credit notes:"
)
flow2 = doc.add_paragraph()
r2 = flow2.add_run(
    "Sales Invoice (Paid/Unpaid)\n"
    "    --> create_sales_invoice_return (qty must be negative)\n"
    "Sales Return (Credit Note)\n"
    "    --> create_payment_entry (to refund or adjust balance)\n"
    "Payment Entry (Submitted) + Outstanding adjusted"
)
r2.font.name = "Consolas"
r2.font.size = Pt(9)

doc.add_paragraph()
doc.add_heading("Workflow 3: SO to DN to SI (Optional - when DN is enabled)", level=2)
doc.add_paragraph(
    "This workflow is used when 'Ignore Create Delivery Note' is unchecked:"
)
flow3 = doc.add_paragraph()
r3 = flow3.add_run(
    "Sales Order (Draft)\n"
    "    --> submit_sales_order\n"
    "Sales Order (Submitted)\n"
    "    --> create_dn_from_so\n"
    "Delivery Note (Draft)\n"
    "    --> submit_delivery_note\n"
    "Delivery Note (Submitted / To Bill)\n"
    "    --> create_si_from_dn\n"
    "Sales Invoice (Draft)\n"
    "    --> submit_sales_invoice\n"
    "Sales Invoice (Unpaid)\n"
    "    --> create_payment_entry\n"
    "Payment Entry (Submitted) + Sales Invoice (Paid)"
)
r3.font.name = "Consolas"
r3.font.size = Pt(9)

# ── SAVE ─────────────────────────────────────────────────────────────────────

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Basket4Me_PWA_API_Reference_V3.docx")
doc.save(output_path)
print(f"Document saved to: {output_path}")
