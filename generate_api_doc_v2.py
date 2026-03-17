#!/usr/bin/env python3
"""
Generate Basket4Me PWA - Van Sales API Reference V2 (Word Document)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import json
import os

# ── Helpers ──────────────────────────────────────────────────────────────────

BASE_URL = "https://devassociate.frappe.cloud"
API_PREFIX = "/api/method/basket4me_pwa.api."
AUTH_PREFIX = "/api/method/basket4me_pwa.auth."

doc = Document()

# Default font
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(10)

# Heading styles
for i in range(1, 4):
    hs = doc.styles[f"Heading {i}"]
    hs.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_text(cell, text, bold=False, color=None, font_name=None, font_size=None, align=None):
    """Set cell text with formatting."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = color
    if font_name:
        run.font.name = font_name
    if font_size:
        run.font.size = font_size

def add_header_row(table, headers, color="2E75B6"):
    """Format the header row of a table."""
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, color)
        set_cell_text(cell, h, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), font_size=Pt(9))

def make_table(headers, rows, header_color="2E75B6"):
    """Create a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_header_row(table, headers, header_color)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            set_cell_text(cell, str(val), font_size=Pt(9))
            if ri % 2 == 1:
                set_cell_shading(cell, "F2F7FB")
    return table

def add_endpoint_summary_table(endpoints):
    """Add a numbered summary table of endpoints."""
    rows = []
    for i, ep in enumerate(endpoints, 1):
        rows.append([str(i), ep["path"], ep["method"], ep["desc"]])
    make_table(["#", "Endpoint Path", "Method", "Description"], rows)

def add_json_block(title, data):
    """Add a JSON code block."""
    p = doc.add_paragraph()
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(10)
    json_str = json.dumps(data, indent=2)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(json_str)
    run2.font.name = "Consolas"
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_param_table(params):
    """Add a parameter table with green header."""
    rows = []
    for p in params:
        rows.append([p["name"], p["type"], p.get("required", "Yes"), p.get("desc", "")])
    make_table(["Parameter", "Type", "Required", "Description"], rows, header_color="2D8B4E")

def add_endpoint_detail(path, method, description, params=None, request_body=None, response_body=None):
    """Add detailed documentation for one endpoint."""
    doc.add_paragraph()
    # Endpoint URL
    p = doc.add_paragraph()
    run_m = p.add_run(f"  {method}  ")
    run_m.bold = True
    run_m.font.size = Pt(10)
    if method == "POST":
        run_m.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # We can't shade a run, so use text color approach
        run_m.font.color.rgb = RGBColor(0xE0, 0x6C, 0x00)
    else:
        run_m.font.color.rgb = RGBColor(0x2D, 0x8B, 0x4E)
    run_url = p.add_run(f"  {BASE_URL}{path}")
    run_url.font.name = "Consolas"
    run_url.font.size = Pt(9)
    run_url.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

    p2 = doc.add_paragraph(description)
    p2.runs[0].font.size = Pt(9)
    p2.runs[0].italic = True

    if params:
        doc.add_paragraph("Parameters:", style="List Bullet")
        add_param_table(params)

    if request_body:
        add_json_block("Request Body:", request_body)

    if response_body:
        add_json_block("Response:", response_body)

    # Thin separator
    p_sep = doc.add_paragraph()
    run_sep = p_sep.add_run("_" * 90)
    run_sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    run_sep.font.size = Pt(6)

def section_heading(num, title):
    doc.add_page_break()
    doc.add_heading(f"Section {num}: {title}", level=1)

# ── TITLE PAGE ───────────────────────────────────────────────────────────────

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title_p.add_run("Basket4Me PWA")
run_t.bold = True
run_t.font.size = Pt(32)
run_t.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = sub_p.add_run("Van Sales Mobile App\nAPI Reference")
run_s.font.size = Pt(22)
run_s.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

doc.add_paragraph()
doc.add_paragraph()

info_p = doc.add_paragraph()
info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for label, value in [
    ("Base URL: ", BASE_URL),
    ("\nVersion: ", "2.0"),
    ("\nDate: ", "March 2026"),
    ("\nAuthorization: ", "Cookie-based (sid + CSRF Token)"),
]:
    r_l = info_p.add_run(label)
    r_l.bold = True
    r_l.font.size = Pt(12)
    r_l.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r_v = info_p.add_run(value)
    r_v.font.size = Pt(12)
    r_v.font.name = "Consolas"
    r_v.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

# ── TABLE OF CONTENTS ────────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading("Table of Contents", level=1)

toc_items = [
    "1. Authentication (3 endpoints)",
    "2. Dashboard (3 endpoints)",
    "3. Customer Management (11 endpoints)",
    "4. Items (6 endpoints)",
    "5. Sales Order (8 endpoints)",
    "6. Delivery Note (10 endpoints)",
    "7. Sales Invoice (10 endpoints)",
    "8. Sales Return (3 endpoints)",
    "9. Receipts / Payments (5 endpoints)",
    "10. Config & Settings (8 endpoints)",
    "11. Print (3 endpoints)",
    "12. Shift Management (3 endpoints)",
    "",
    "Appendix A: Standard Error Response Format",
    "Appendix B: HTTP Status Codes",
    "Appendix C: Workflow Diagrams",
]
for item in toc_items:
    if item == "":
        doc.add_paragraph()
        continue
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].font.size = Pt(11)

# ── AUTH HEADERS SECTION ─────────────────────────────────────────────────────

doc.add_page_break()
doc.add_heading("Authentication & Authorization", level=1)

doc.add_paragraph(
    "The Basket4Me PWA API uses cookie-based authentication. "
    "The Login endpoint returns a session ID (sid) and a CSRF token. "
    "These must be included in all subsequent requests as described below."
)

doc.add_heading("How It Works", level=2)
doc.add_paragraph(
    "1. Call the Login endpoint with usr and pwd.\n"
    "2. The response contains sid and csrf_token.\n"
    "3. Store both values on the client for all future requests."
)

doc.add_heading("Required Headers by Method", level=2)

make_table(
    ["HTTP Method", "Required Headers", "Example"],
    [
        ["GET", "Cookie: sid=<session_id>", "Cookie: sid=abc123xyz"],
        ["POST", "Cookie: sid=<session_id>\nX-Frappe-CSRF-Token: <token>",
         "Cookie: sid=abc123xyz\nX-Frappe-CSRF-Token: tok456"],
    ],
)

doc.add_paragraph()
p_note = doc.add_paragraph()
run_note = p_note.add_run("Note: ")
run_note.bold = True
p_note.add_run(
    "All POST requests MUST include the X-Frappe-CSRF-Token header in addition to the Cookie header. "
    "GET requests only need the Cookie header. Omitting these headers will result in a 403 Forbidden error."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1: AUTHENTICATION
# ══════════════════════════════════════════════════════════════════════════════

section_heading(1, "Authentication")

auth_endpoints = [
    {"path": AUTH_PREFIX + "user_login", "method": "POST", "desc": "Login and obtain session credentials"},
    {"path": AUTH_PREFIX + "logout", "method": "POST", "desc": "Logout and invalidate session"},
    {"path": AUTH_PREFIX + "get_user_details", "method": "POST", "desc": "Get detailed user profile"},
]
add_endpoint_summary_table(auth_endpoints)

doc.add_heading("1.1 User Login", level=2)
add_endpoint_detail(
    AUTH_PREFIX + "user_login", "POST",
    "Authenticate a user and obtain session credentials for subsequent API calls.",
    params=[
        {"name": "usr", "type": "String", "required": "Yes", "desc": "Username or email address"},
        {"name": "pwd", "type": "String", "required": "Yes", "desc": "User password"},
    ],
    request_body={"usr": "salesman@example.com", "pwd": "securepassword123"},
    response_body={
        "message": {
            "sid": "abc123sessionid",
            "csrf_token": "tok456csrftoken",
            "api_key": "api_key_value",
            "api_secret": "api_secret_value",
            "username": "salesman@example.com",
            "email": "salesman@example.com",
            "full_name": "John Doe",
            "company_name": "My Trading LLC",
            "company_logo": "/files/logo.png",
            "user_image": "/files/user.png"
        }
    },
)

doc.add_heading("1.2 Logout", level=2)
add_endpoint_detail(
    AUTH_PREFIX + "logout", "POST",
    "Invalidate the current session and logout the user.",
    params=[{"name": "usr", "type": "String", "required": "Yes", "desc": "Username or email"}],
    request_body={"usr": "salesman@example.com"},
    response_body={"message": "Logged Out"},
)

doc.add_heading("1.3 Get User Details", level=2)
add_endpoint_detail(
    AUTH_PREFIX + "get_user_details", "POST",
    "Retrieve detailed profile information for a user.",
    params=[{"name": "user_id", "type": "String", "required": "Yes", "desc": "User ID or email"}],
    request_body={"user_id": "salesman@example.com"},
    response_body={
        "message": {
            "username": "salesman@example.com",
            "full_name": "John Doe",
            "email": "salesman@example.com",
            "company_name": "My Trading LLC",
            "company_logo": "/files/logo.png",
            "user_image": "/files/user.png"
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2: DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════

section_heading(2, "Dashboard")

dash_endpoints = [
    {"path": API_PREFIX + "get_dashboard_summary", "method": "GET", "desc": "Dashboard summary with orders, invoices, returns, collections"},
    {"path": API_PREFIX + "get_sales_metrics", "method": "GET", "desc": "Sales metrics and KPIs"},
    {"path": API_PREFIX + "get_daily_sales_report", "method": "GET", "desc": "Daily sales breakdown by mode of payment"},
]
add_endpoint_summary_table(dash_endpoints)

doc.add_heading("2.1 Get Dashboard Summary", level=2)
add_endpoint_detail(
    API_PREFIX + "get_dashboard_summary", "GET",
    "Retrieve a consolidated dashboard summary for orders, invoices, returns, and collections.",
    params=[
        {"name": "period", "type": "String", "required": "No", "desc": "Period filter: daily, weekly, monthly"},
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date (YYYY-MM-DD)"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date (YYYY-MM-DD)"},
    ],
    response_body={
        "message": {
            "orders": {"count": 25, "total_amount": 15000.00},
            "invoices": {"count": 20, "total_amount": 12500.00},
            "returns": {"count": 3, "total_amount": 1200.00},
            "collections": {"count": 18, "total_amount": 11000.00},
            "visited_customers": 15
        }
    },
)

doc.add_heading("2.2 Get Sales Metrics", level=2)
add_endpoint_detail(
    API_PREFIX + "get_sales_metrics", "GET",
    "Retrieve key sales performance metrics for the salesperson.",
    params=[
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date"},
        {"name": "sales_person_filter", "type": "String", "required": "No", "desc": "Filter by specific salesperson"},
    ],
    response_body={
        "message": {
            "total_sales_amount": 45000.00,
            "today_sales_count": 8,
            "today_sales_amount": 3200.00,
            "total_collections": 38000.00,
            "outstanding_amount": 7000.00
        }
    },
)

doc.add_heading("2.3 Get Daily Sales Report", level=2)
add_endpoint_detail(
    API_PREFIX + "get_daily_sales_report", "GET",
    "Retrieve daily sales breakdown by mode of payment (Cash, Card, Credit, etc.).",
    response_body={
        "message": [
            {"date": "2026-03-16", "cash": 5000.00, "card": 2000.00, "credit": 3000.00, "total": 10000.00},
            {"date": "2026-03-15", "cash": 4500.00, "card": 1800.00, "credit": 2700.00, "total": 9000.00}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3: CUSTOMER MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

section_heading(3, "Customer Management")

cust_endpoints = [
    {"path": API_PREFIX + "get_customer_list_v2", "method": "GET", "desc": "Advanced customer list with filters, visit status, outstanding balance"},
    {"path": API_PREFIX + "get_customer_list", "method": "GET", "desc": "Simple customer list with name/mobile search"},
    {"path": API_PREFIX + "get_customer_detail", "method": "POST", "desc": "Full customer details with addresses and credit info"},
    {"path": API_PREFIX + "create_customer", "method": "POST", "desc": "Create a new customer"},
    {"path": API_PREFIX + "get_customer_invoices", "method": "GET", "desc": "List invoices for a customer"},
    {"path": API_PREFIX + "customer_invoice_details", "method": "POST", "desc": "Details of a specific customer invoice"},
    {"path": API_PREFIX + "get_customer_invoice_aging", "method": "GET", "desc": "Invoice aging report for a customer"},
    {"path": API_PREFIX + "get_customer_balance_summary", "method": "GET", "desc": "Customer balance summary (old, current, total, credit limit)"},
    {"path": API_PREFIX + "mark_customer_visit", "method": "POST", "desc": "Mark a customer visit with GPS location"},
    {"path": API_PREFIX + "get_customer_visits", "method": "GET", "desc": "List customer visits within date range"},
    {"path": API_PREFIX + "get_territories", "method": "GET", "desc": "List of all routes/territories"},
]
add_endpoint_summary_table(cust_endpoints)

doc.add_heading("3.1 Get Customer List V2", level=2)
add_endpoint_detail(
    API_PREFIX + "get_customer_list_v2", "GET",
    "Advanced customer listing with filters for territory, route, visit status, and pagination.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by customer name"},
        {"name": "mobile_no", "type": "String", "required": "No", "desc": "Filter by mobile number"},
        {"name": "territory", "type": "String", "required": "No", "desc": "Filter by territory"},
        {"name": "route", "type": "String", "required": "No", "desc": "Filter by route"},
        {"name": "visit_status", "type": "String", "required": "No", "desc": "Filter by visit status (visited/not_visited)"},
        {"name": "page_number", "type": "Integer", "required": "No", "desc": "Page number for pagination"},
        {"name": "page_size", "type": "Integer", "required": "No", "desc": "Records per page (default 20)"},
    ],
    response_body={
        "message": {
            "customers": [
                {
                    "name": "CUST-0001",
                    "customer_name": "ABC Trading",
                    "mobile_no": "+971501234567",
                    "territory": "Dubai",
                    "outstanding_balance": 5000.00,
                    "visit_status": "visited"
                }
            ],
            "total_count": 150
        }
    },
)

doc.add_heading("3.2 Create Customer", level=2)
add_endpoint_detail(
    API_PREFIX + "create_customer", "POST",
    "Create a new customer record with address and contact details.",
    params=[
        {"name": "params.customerdetails.customer_name", "type": "String", "required": "Yes", "desc": "Customer name"},
        {"name": "params.customerdetails.customer_group", "type": "String", "required": "Yes", "desc": "Customer group"},
        {"name": "params.customerdetails.territory", "type": "String", "required": "No", "desc": "Territory/route"},
        {"name": "params.customerdetails.address_line1", "type": "String", "required": "No", "desc": "Address line 1"},
        {"name": "params.customerdetails.city", "type": "String", "required": "No", "desc": "City"},
        {"name": "params.customerdetails.mobile_no", "type": "String", "required": "No", "desc": "Mobile number"},
        {"name": "params.customerdetails.email_id", "type": "String", "required": "No", "desc": "Email address"},
    ],
    request_body={
        "params": {
            "customerdetails": {
                "customer_name": "New Shop LLC",
                "customer_group": "Retail",
                "territory": "Dubai",
                "address_line1": "Shop 5, Al Karama",
                "city": "Dubai",
                "mobile_no": "+971501234567",
                "email_id": "newshop@example.com"
            }
        }
    },
    response_body={
        "message": {
            "name": "CUST-0045",
            "customer_name": "New Shop LLC",
            "customer_group": "Retail",
            "territory": "Dubai"
        }
    },
)

doc.add_heading("3.3 Mark Customer Visit", level=2)
add_endpoint_detail(
    API_PREFIX + "mark_customer_visit", "POST",
    "Record a customer visit with GPS coordinates and optional remarks.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
        {"name": "latitude", "type": "Float", "required": "Yes", "desc": "GPS latitude"},
        {"name": "longitude", "type": "Float", "required": "Yes", "desc": "GPS longitude"},
        {"name": "remarks", "type": "String", "required": "No", "desc": "Visit remarks/notes"},
    ],
    request_body={
        "customer": "CUST-0001",
        "latitude": 25.2048,
        "longitude": 55.2708,
        "remarks": "Collected payment, discussed new order"
    },
    response_body={"message": {"status": "success", "visit_id": "VISIT-0012"}},
)

doc.add_heading("3.4 Get Customer Balance Summary", level=2)
add_endpoint_detail(
    API_PREFIX + "get_customer_balance_summary", "GET",
    "Retrieve customer balance details including old balance, current balance, total, and credit limit.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    response_body={
        "message": {
            "old_balance": 3000.00,
            "current_balance": 2000.00,
            "total_balance": 5000.00,
            "credit_limit": 10000.00
        }
    },
)

doc.add_heading("3.5 Get Customer Visits", level=2)
add_endpoint_detail(
    API_PREFIX + "get_customer_visits", "GET",
    "Retrieve list of customer visits within a date range.",
    params=[
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date (YYYY-MM-DD)"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date (YYYY-MM-DD)"},
    ],
    response_body={
        "message": [
            {"visit_id": "VISIT-0012", "customer": "CUST-0001", "customer_name": "ABC Trading",
             "latitude": 25.2048, "longitude": 55.2708, "visit_time": "2026-03-16 10:30:00",
             "remarks": "Collected payment"}
        ]
    },
)

doc.add_heading("3.6 Get Territories", level=2)
add_endpoint_detail(
    API_PREFIX + "get_territories", "GET",
    "Retrieve list of all available routes and territories.",
    response_body={
        "message": [
            {"name": "Dubai", "parent_territory": "UAE"},
            {"name": "Abu Dhabi", "parent_territory": "UAE"},
            {"name": "Sharjah", "parent_territory": "UAE"}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4: ITEMS
# ══════════════════════════════════════════════════════════════════════════════

section_heading(4, "Items")

item_endpoints = [
    {"path": API_PREFIX + "get_item_list", "method": "GET", "desc": "List items with pricing, supports pagination"},
    {"path": API_PREFIX + "get_item_detail", "method": "POST", "desc": "Full item details"},
    {"path": API_PREFIX + "get_item_by_barcode", "method": "GET", "desc": "Lookup item by barcode"},
    {"path": API_PREFIX + "get_item_stock", "method": "GET", "desc": "Stock levels per warehouse"},
    {"path": API_PREFIX + "get_item_pricing_details", "method": "GET", "desc": "Pricing details including MRP, UOM rates, tax"},
    {"path": API_PREFIX + "get_item_uom_details", "method": "GET", "desc": "All UOMs with conversion factors for an item"},
]
add_endpoint_summary_table(item_endpoints)

doc.add_heading("4.1 Get Item List", level=2)
add_endpoint_detail(
    API_PREFIX + "get_item_list", "GET",
    "Retrieve paginated list of items with pricing information.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by item code"},
        {"name": "item_name", "type": "String", "required": "No", "desc": "Filter by item name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Customer for customer-specific pricing"},
        {"name": "page_number", "type": "Integer", "required": "No", "desc": "Page number"},
        {"name": "page_size", "type": "Integer", "required": "No", "desc": "Records per page"},
    ],
    response_body={
        "message": {
            "items": [
                {
                    "item_code": "ITEM-001",
                    "item_name": "Coca Cola 330ml",
                    "item_group": "Beverages",
                    "stock_uom": "Nos",
                    "rate": 2.50,
                    "image": "/files/coca-cola.png"
                }
            ],
            "total_count": 500
        }
    },
)

doc.add_heading("4.2 Get Item by Barcode", level=2)
add_endpoint_detail(
    API_PREFIX + "get_item_by_barcode", "GET",
    "Look up an item using its barcode.",
    params=[
        {"name": "barcode", "type": "String", "required": "Yes", "desc": "Barcode value to look up"},
    ],
    response_body={
        "message": {
            "item_code": "ITEM-001",
            "item_name": "Coca Cola 330ml",
            "barcode": "6291001000012",
            "stock_uom": "Nos",
            "rate": 2.50
        }
    },
)

doc.add_heading("4.3 Get Item Stock", level=2)
add_endpoint_detail(
    API_PREFIX + "get_item_stock", "GET",
    "Retrieve stock levels for an item across warehouses.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "warehouse", "type": "String", "required": "No", "desc": "Specific warehouse (optional)"},
    ],
    response_body={
        "message": [
            {
                "warehouse": "Van 1 - MT",
                "actual_qty": 100,
                "reserved_qty": 10,
                "projected_qty": 90
            },
            {
                "warehouse": "Main Store - MT",
                "actual_qty": 500,
                "reserved_qty": 50,
                "projected_qty": 450
            }
        ]
    },
)

doc.add_heading("4.4 Get Item Pricing Details", level=2)
add_endpoint_detail(
    API_PREFIX + "get_item_pricing_details", "GET",
    "Retrieve comprehensive pricing details for an item, optionally for a specific customer and UOM.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Customer for customer-specific pricing"},
        {"name": "uom", "type": "String", "required": "No", "desc": "Unit of measure"},
    ],
    response_body={
        "message": {
            "mrp": 3.00,
            "last_sales_price": 2.50,
            "price_list_rate": 2.75,
            "uoms": [
                {"uom": "Nos", "conversion_factor": 1},
                {"uom": "Box", "conversion_factor": 24},
                {"uom": "Carton", "conversion_factor": 48}
            ],
            "tax_rate": 5.0
        }
    },
)

doc.add_heading("4.5 Get Item UOM Details", level=2)
add_endpoint_detail(
    API_PREFIX + "get_item_uom_details", "GET",
    "Retrieve all available UOMs with conversion factors for an item.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
    ],
    response_body={
        "message": [
            {"uom": "Nos", "conversion_factor": 1},
            {"uom": "Box", "conversion_factor": 24},
            {"uom": "Carton", "conversion_factor": 48}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5: SALES ORDER
# ══════════════════════════════════════════════════════════════════════════════

section_heading(5, "Sales Order")

so_endpoints = [
    {"path": API_PREFIX + "create_sales_order", "method": "POST", "desc": "Create a new sales order"},
    {"path": API_PREFIX + "update_sales_order", "method": "POST", "desc": "Update an existing draft sales order"},
    {"path": API_PREFIX + "submit_sales_order", "method": "POST", "desc": "Submit a draft sales order"},
    {"path": API_PREFIX + "cancel_sales_order", "method": "POST", "desc": "Cancel a submitted sales order"},
    {"path": API_PREFIX + "delete_sales_order", "method": "POST", "desc": "Delete a draft sales order"},
    {"path": API_PREFIX + "get_sales_order_list", "method": "GET", "desc": "List sales orders with filters and pagination"},
    {"path": API_PREFIX + "get_sales_order_detail", "method": "GET", "desc": "Get full details of a sales order"},
    {"path": API_PREFIX + "convert_so_to_si", "method": "POST", "desc": "Convert sales order(s) to sales invoice with payments"},
]
add_endpoint_summary_table(so_endpoints)

doc.add_heading("5.1 Create Sales Order", level=2)
add_endpoint_detail(
    API_PREFIX + "create_sales_order", "POST",
    "Create a new sales order with line items and optional payment type.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
        {"name": "items", "type": "Array", "required": "Yes", "desc": "Array of item objects"},
        {"name": "items[].item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "items[].qty", "type": "Float", "required": "Yes", "desc": "Quantity"},
        {"name": "items[].uom", "type": "String", "required": "No", "desc": "Unit of measure"},
        {"name": "items[].rate", "type": "Float", "required": "No", "desc": "Price per unit"},
        {"name": "items[].discount_percentage", "type": "Float", "required": "No", "desc": "Discount percentage"},
        {"name": "items[].discount_amount", "type": "Float", "required": "No", "desc": "Discount amount"},
        {"name": "items[].is_free_item", "type": "Boolean", "required": "No", "desc": "Whether item is free (0/1)"},
        {"name": "items[].batch_no", "type": "String", "required": "No", "desc": "Batch number"},
        {"name": "delivery_date", "type": "Date", "required": "No", "desc": "Expected delivery date"},
        {"name": "po_no", "type": "String", "required": "No", "desc": "Customer purchase order number"},
        {"name": "remarks", "type": "String", "required": "No", "desc": "Order remarks"},
        {"name": "custom_payment_type", "type": "String", "required": "No", "desc": "Payment type (Cash/Credit)"},
    ],
    request_body={
        "customer": "CUST-0001",
        "delivery_date": "2026-03-20",
        "po_no": "PO-2026-100",
        "remarks": "Urgent delivery required",
        "custom_payment_type": "Cash",
        "items": [
            {
                "item_code": "ITEM-001",
                "qty": 24,
                "uom": "Nos",
                "rate": 2.50,
                "discount_percentage": 0,
                "discount_amount": 0,
                "is_free_item": 0,
                "batch_no": "BATCH-001"
            },
            {
                "item_code": "ITEM-002",
                "qty": 1,
                "uom": "Box",
                "rate": 50.00,
                "is_free_item": 0
            }
        ]
    },
    response_body={
        "message": {
            "name": "SO-2026-00045",
            "status": "Draft",
            "customer": "CUST-0001",
            "grand_total": 110.00,
            "delivery_date": "2026-03-20"
        }
    },
)

doc.add_heading("5.2 Get Sales Order List", level=2)
add_endpoint_detail(
    API_PREFIX + "get_sales_order_list", "GET",
    "Retrieve a filtered and paginated list of sales orders.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by SO name/ID"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status (Draft/Submitted/Cancelled)"},
        {"name": "search", "type": "String", "required": "No", "desc": "General search text"},
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date filter"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date filter"},
        {"name": "page_number", "type": "Integer", "required": "No", "desc": "Page number"},
        {"name": "page_size", "type": "Integer", "required": "No", "desc": "Records per page"},
    ],
    response_body={
        "message": {
            "sales_orders": [
                {
                    "name": "SO-2026-00045",
                    "customer": "CUST-0001",
                    "customer_name": "ABC Trading",
                    "status": "Draft",
                    "grand_total": 110.00,
                    "transaction_date": "2026-03-16"
                }
            ],
            "total_count": 45
        }
    },
)

doc.add_heading("5.3 Get Sales Order Detail", level=2)
add_endpoint_detail(
    API_PREFIX + "get_sales_order_detail", "GET",
    "Get full details of a single Sales Order including items, taxes, and sales team.",
    params=[
        {"name": "name", "type": "String", "required": "Yes", "desc": "Sales Order ID (e.g. SO-2026-00045)"},
    ],
    request_body=None,
    response_body={
        "message": "Sales Order fetched successfully",
        "data": {
            "name": "SO-2026-00045",
            "customer": "CUST-0001",
            "customer_name": "Vimal Store Cherthala",
            "transaction_date": "2026-03-17",
            "delivery_date": "2026-03-20",
            "docstatus": 1,
            "status": "To Deliver and Bill",
            "company": "Dev Associate",
            "currency": "INR",
            "selling_price_list": "Standard Selling",
            "total": 213.60,
            "net_total": 213.60,
            "total_taxes_and_charges": 0,
            "grand_total": 213.60,
            "per_delivered": 0,
            "per_billed": 0,
            "po_no": None,
            "items": [
                {
                    "name": "row-id-001",
                    "item_code": "AF00890",
                    "item_name": "B Unibic Cashew Rs.10*192",
                    "description": "B Unibic Cashew Rs.10*192",
                    "qty": 2,
                    "uom": "Nos",
                    "rate": 106.80,
                    "price_list_rate": 106.80,
                    "discount_percentage": 0,
                    "discount_amount": 0,
                    "amount": 213.60,
                    "warehouse": "Stores - DA",
                    "delivery_date": "2026-03-20"
                }
            ],
            "taxes": [],
            "sales_team": [
                {
                    "sales_person": "aswathy",
                    "allocated_percentage": 100
                }
            ]
        },
        "success": True
    },
)

doc.add_heading("5.4 Convert Sales Order to Sales Invoice", level=2)
add_endpoint_detail(
    API_PREFIX + "convert_so_to_si", "POST",
    "Convert one or more sales orders into a sales invoice with payment entries.",
    params=[
        {"name": "sales_orders", "type": "Array", "required": "Yes", "desc": "List of SO names to convert"},
        {"name": "payments", "type": "Array", "required": "No", "desc": "Payment details for the invoice"},
    ],
    request_body={
        "sales_orders": ["SO-2026-00045", "SO-2026-00046"],
        "payments": [
            {"mode_of_payment": "Cash", "amount": 200.00},
            {"mode_of_payment": "Credit", "amount": 50.00}
        ]
    },
    response_body={
        "message": {
            "sales_invoice": "SINV-2026-00030",
            "status": "Draft",
            "grand_total": 250.00
        }
    },
)

doc.add_heading("5.5 Submit / Cancel / Delete Sales Order", level=2)
doc.add_paragraph(
    "These endpoints accept the sales order name and perform the respective action."
)
make_table(
    ["Endpoint", "Method", "Parameter", "Description"],
    [
        [API_PREFIX + "submit_sales_order", "POST", "name (String)", "Submit a draft SO"],
        [API_PREFIX + "cancel_sales_order", "POST", "name (String)", "Cancel a submitted SO"],
        [API_PREFIX + "delete_sales_order", "POST", "name (String)", "Permanently delete a draft SO"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6: DELIVERY NOTE
# ══════════════════════════════════════════════════════════════════════════════

section_heading(6, "Delivery Note")

dn_endpoints = [
    {"path": API_PREFIX + "create_delivery_note", "method": "POST", "desc": "Create a new delivery note"},
    {"path": API_PREFIX + "update_delivery_note", "method": "POST", "desc": "Update a draft delivery note"},
    {"path": API_PREFIX + "submit_delivery_note", "method": "POST", "desc": "Submit a draft delivery note"},
    {"path": API_PREFIX + "cancel_delivery_note", "method": "POST", "desc": "Cancel a submitted delivery note"},
    {"path": API_PREFIX + "delete_delivery_note", "method": "POST", "desc": "Delete a draft delivery note"},
    {"path": API_PREFIX + "get_delivery_note_list", "method": "GET", "desc": "List delivery notes with filters"},
    {"path": API_PREFIX + "get_delivery_note_detail", "method": "GET", "desc": "Get full details of a delivery note"},
    {"path": API_PREFIX + "get_pending_so_for_dn", "method": "GET", "desc": "Get pending SOs available for DN creation"},
    {"path": API_PREFIX + "create_dn_from_so", "method": "POST", "desc": "Create DN from one or more sales orders"},
    {"path": API_PREFIX + "create_si_from_dn", "method": "POST", "desc": "Create sales invoice from delivery note(s)"},
]
add_endpoint_summary_table(dn_endpoints)

doc.add_heading("6.1 Create Delivery Note", level=2)
add_endpoint_detail(
    API_PREFIX + "create_delivery_note", "POST",
    "Create a new delivery note with line items.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
        {"name": "items", "type": "Array", "required": "Yes", "desc": "Array of item objects with item_code, qty, uom, rate"},
        {"name": "posting_date", "type": "Date", "required": "No", "desc": "Posting date"},
    ],
    request_body={
        "customer": "CUST-0001",
        "posting_date": "2026-03-16",
        "items": [
            {"item_code": "ITEM-001", "qty": 24, "uom": "Nos", "rate": 2.50},
            {"item_code": "ITEM-002", "qty": 1, "uom": "Box", "rate": 50.00}
        ]
    },
    response_body={
        "message": {
            "name": "DN-2026-00020",
            "status": "Draft",
            "customer": "CUST-0001",
            "grand_total": 110.00
        }
    },
)

doc.add_heading("6.2 Get Delivery Note List", level=2)
add_endpoint_detail(
    API_PREFIX + "get_delivery_note_list", "GET",
    "Retrieve filtered list of delivery notes.",
    params=[
        {"name": "name", "type": "String", "required": "No", "desc": "Filter by DN name"},
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status"},
        {"name": "search", "type": "String", "required": "No", "desc": "Search text"},
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date"},
    ],
    response_body={
        "message": {
            "delivery_notes": [
                {
                    "name": "DN-2026-00020",
                    "customer": "CUST-0001",
                    "customer_name": "ABC Trading",
                    "status": "Draft",
                    "grand_total": 110.00,
                    "posting_date": "2026-03-16"
                }
            ],
            "total_count": 20
        }
    },
)

doc.add_heading("6.3 Get Pending SO for DN", level=2)
add_endpoint_detail(
    API_PREFIX + "get_pending_so_for_dn", "GET",
    "Retrieve sales orders pending delivery for a customer.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    response_body={
        "message": [
            {
                "name": "SO-2026-00045",
                "customer_name": "ABC Trading",
                "grand_total": 110.00,
                "delivery_date": "2026-03-20",
                "items": [
                    {"item_code": "ITEM-001", "qty": 24, "delivered_qty": 0, "pending_qty": 24}
                ]
            }
        ]
    },
)

doc.add_heading("6.4 Create DN from Sales Order", level=2)
add_endpoint_detail(
    API_PREFIX + "create_dn_from_so", "POST",
    "Create a delivery note from one or more submitted sales orders.",
    params=[
        {"name": "sales_orders", "type": "Array", "required": "Yes", "desc": "List of SO names"},
    ],
    request_body={"sales_orders": ["SO-2026-00045"]},
    response_body={
        "message": {
            "name": "DN-2026-00021",
            "status": "Draft",
            "grand_total": 110.00
        }
    },
)

doc.add_heading("6.5 Create Sales Invoice from DN", level=2)
add_endpoint_detail(
    API_PREFIX + "create_si_from_dn", "POST",
    "Create a sales invoice from one or more submitted delivery notes with optional payments.",
    params=[
        {"name": "delivery_notes", "type": "Array", "required": "Yes", "desc": "List of DN names"},
        {"name": "payments", "type": "Array", "required": "No", "desc": "Payment details"},
    ],
    request_body={
        "delivery_notes": ["DN-2026-00020"],
        "payments": [{"mode_of_payment": "Cash", "amount": 110.00}]
    },
    response_body={
        "message": {
            "sales_invoice": "SINV-2026-00031",
            "status": "Draft",
            "grand_total": 110.00
        }
    },
)

doc.add_heading("6.6 Submit / Cancel / Delete Delivery Note", level=2)
make_table(
    ["Endpoint", "Method", "Parameter", "Description"],
    [
        [API_PREFIX + "submit_delivery_note", "POST", "name (String)", "Submit a draft DN"],
        [API_PREFIX + "cancel_delivery_note", "POST", "name (String)", "Cancel a submitted DN"],
        [API_PREFIX + "delete_delivery_note", "POST", "name (String)", "Permanently delete a draft DN"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7: SALES INVOICE
# ══════════════════════════════════════════════════════════════════════════════

section_heading(7, "Sales Invoice")

si_endpoints = [
    {"path": API_PREFIX + "create_sales_invoice", "method": "POST", "desc": "Create a new sales invoice"},
    {"path": API_PREFIX + "update_sales_invoice", "method": "POST", "desc": "Update a draft sales invoice"},
    {"path": API_PREFIX + "submit_sales_invoice", "method": "POST", "desc": "Submit a draft sales invoice"},
    {"path": API_PREFIX + "cancel_sales_invoice", "method": "POST", "desc": "Cancel a submitted sales invoice"},
    {"path": API_PREFIX + "delete_sales_invoice", "method": "POST", "desc": "Delete a draft sales invoice"},
    {"path": API_PREFIX + "get_invoice_list", "method": "GET", "desc": "List sales invoices"},
    {"path": API_PREFIX + "get_invoice_detail", "method": "GET", "desc": "Get invoice details (GET)"},
    {"path": API_PREFIX + "invoice_details", "method": "POST", "desc": "Get invoice details (POST)"},
    {"path": API_PREFIX + "get_uninvoiced_dn", "method": "GET", "desc": "Get delivery notes not yet invoiced"},
    {"path": API_PREFIX + "get_invoices_for_return", "method": "GET", "desc": "Get invoices eligible for return"},
]
add_endpoint_summary_table(si_endpoints)

doc.add_heading("7.1 Create Sales Invoice", level=2)
add_endpoint_detail(
    API_PREFIX + "create_sales_invoice", "POST",
    "Create a new sales invoice with line items and optional payments.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
        {"name": "items", "type": "Array", "required": "Yes", "desc": "Array of item objects"},
        {"name": "payments", "type": "Array", "required": "No", "desc": "Payment split details"},
        {"name": "posting_date", "type": "Date", "required": "No", "desc": "Invoice posting date"},
        {"name": "due_date", "type": "Date", "required": "No", "desc": "Payment due date"},
    ],
    request_body={
        "customer": "CUST-0001",
        "posting_date": "2026-03-16",
        "due_date": "2026-04-15",
        "items": [
            {"item_code": "ITEM-001", "qty": 24, "uom": "Nos", "rate": 2.50},
            {"item_code": "ITEM-002", "qty": 1, "uom": "Box", "rate": 50.00}
        ],
        "payments": [
            {"mode_of_payment": "Cash", "amount": 60.00},
            {"mode_of_payment": "Credit", "amount": 50.00}
        ]
    },
    response_body={
        "message": {
            "name": "SINV-2026-00032",
            "status": "Draft",
            "customer": "CUST-0001",
            "grand_total": 110.00
        }
    },
)

doc.add_heading("7.2 Get Invoice List", level=2)
add_endpoint_detail(
    API_PREFIX + "get_invoice_list", "GET",
    "Retrieve a list of sales invoices with optional filters.",
    params=[
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "status", "type": "String", "required": "No", "desc": "Filter by status"},
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date"},
        {"name": "page_number", "type": "Integer", "required": "No", "desc": "Page number"},
        {"name": "page_size", "type": "Integer", "required": "No", "desc": "Records per page"},
    ],
    response_body={
        "message": {
            "invoices": [
                {
                    "name": "SINV-2026-00032",
                    "customer": "CUST-0001",
                    "customer_name": "ABC Trading",
                    "status": "Draft",
                    "grand_total": 110.00,
                    "posting_date": "2026-03-16",
                    "outstanding_amount": 50.00
                }
            ],
            "total_count": 32
        }
    },
)

doc.add_heading("7.3 Get Uninvoiced Delivery Notes", level=2)
add_endpoint_detail(
    API_PREFIX + "get_uninvoiced_dn", "GET",
    "Get delivery notes for a customer that have not yet been invoiced.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    response_body={
        "message": [
            {
                "name": "DN-2026-00020",
                "customer_name": "ABC Trading",
                "grand_total": 110.00,
                "posting_date": "2026-03-16",
                "items": [
                    {"item_code": "ITEM-001", "qty": 24, "billed_qty": 0, "pending_qty": 24}
                ]
            }
        ]
    },
)

doc.add_heading("7.4 Get Invoices for Return", level=2)
add_endpoint_detail(
    API_PREFIX + "get_invoices_for_return", "GET",
    "Get invoices eligible for return for a customer, including item-level returned and returnable quantities.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
    ],
    response_body={
        "message": [
            {
                "name": "SINV-2026-00032",
                "customer_name": "ABC Trading",
                "grand_total": 110.00,
                "posting_date": "2026-03-16",
                "items": [
                    {
                        "item_code": "ITEM-001",
                        "item_name": "Coca Cola 330ml",
                        "qty": 24,
                        "returned_qty": 0,
                        "returnable_qty": 24
                    }
                ]
            }
        ]
    },
)

doc.add_heading("7.5 Submit / Cancel / Delete Sales Invoice", level=2)
make_table(
    ["Endpoint", "Method", "Parameter", "Description"],
    [
        [API_PREFIX + "submit_sales_invoice", "POST", "name (String)", "Submit a draft invoice"],
        [API_PREFIX + "cancel_sales_invoice", "POST", "name (String)", "Cancel a submitted invoice"],
        [API_PREFIX + "delete_sales_invoice", "POST", "name (String)", "Permanently delete a draft invoice"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8: SALES RETURN
# ══════════════════════════════════════════════════════════════════════════════

section_heading(8, "Sales Return")

ret_endpoints = [
    {"path": API_PREFIX + "create_sales_invoice_return", "method": "POST", "desc": "Create a return (credit note) against an invoice"},
    {"path": API_PREFIX + "get_return_invoice_list", "method": "GET", "desc": "List return invoices"},
    {"path": API_PREFIX + "get_returned_qty", "method": "GET", "desc": "Get returned quantities for an invoice"},
]
add_endpoint_summary_table(ret_endpoints)

doc.add_heading("8.1 Create Sales Invoice Return", level=2)
add_endpoint_detail(
    API_PREFIX + "create_sales_invoice_return", "POST",
    "Create a return (credit note) against an existing submitted sales invoice. Quantities must be negative.",
    params=[
        {"name": "return_against", "type": "String", "required": "Yes", "desc": "Original invoice name to return against"},
        {"name": "items", "type": "Array", "required": "Yes", "desc": "Items being returned (qty must be negative)"},
        {"name": "items[].item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "items[].qty", "type": "Float", "required": "Yes", "desc": "Quantity (NEGATIVE value)"},
        {"name": "items[].uom", "type": "String", "required": "No", "desc": "Unit of measure"},
    ],
    request_body={
        "return_against": "SINV-2026-00032",
        "items": [
            {"item_code": "ITEM-001", "qty": -5, "uom": "Nos"},
            {"item_code": "ITEM-002", "qty": -1, "uom": "Box"}
        ]
    },
    response_body={
        "message": {
            "name": "SINV-RET-2026-00005",
            "status": "Draft",
            "return_against": "SINV-2026-00032",
            "grand_total": -62.50
        }
    },
)

doc.add_heading("8.2 Get Return Invoice List", level=2)
add_endpoint_detail(
    API_PREFIX + "get_return_invoice_list", "GET",
    "Retrieve a list of return invoices (credit notes).",
    params=[
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "from_date", "type": "Date", "required": "No", "desc": "Start date"},
        {"name": "to_date", "type": "Date", "required": "No", "desc": "End date"},
    ],
    response_body={
        "message": [
            {
                "name": "SINV-RET-2026-00005",
                "customer": "CUST-0001",
                "customer_name": "ABC Trading",
                "return_against": "SINV-2026-00032",
                "grand_total": -62.50,
                "posting_date": "2026-03-16"
            }
        ]
    },
)

doc.add_heading("8.3 Get Returned Qty", level=2)
add_endpoint_detail(
    API_PREFIX + "get_returned_qty", "GET",
    "Get the quantities already returned for a specific invoice.",
    params=[
        {"name": "invoice_name", "type": "String", "required": "Yes", "desc": "Invoice name to check"},
    ],
    response_body={
        "message": [
            {"item_code": "ITEM-001", "qty": 24, "returned_qty": 5, "returnable_qty": 19},
            {"item_code": "ITEM-002", "qty": 1, "returned_qty": 1, "returnable_qty": 0}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 9: RECEIPTS / PAYMENTS
# ══════════════════════════════════════════════════════════════════════════════

section_heading(9, "Receipts / Payments")

pe_endpoints = [
    {"path": API_PREFIX + "create_payment_entry", "method": "POST", "desc": "Create a payment entry / receipt"},
    {"path": API_PREFIX + "cancel_payment_entry", "method": "POST", "desc": "Cancel a payment entry"},
    {"path": API_PREFIX + "get_receipt_list", "method": "GET", "desc": "List receipts/payments with pagination"},
    {"path": API_PREFIX + "receipt_details", "method": "POST", "desc": "Get receipt/payment details"},
    {"path": API_PREFIX + "get_available_modes_of_payment", "method": "GET", "desc": "List available modes of payment"},
]
add_endpoint_summary_table(pe_endpoints)

doc.add_heading("9.1 Create Payment Entry", level=2)
add_endpoint_detail(
    API_PREFIX + "create_payment_entry", "POST",
    "Create a payment entry (receipt) for a customer, optionally allocated against specific invoices.",
    params=[
        {"name": "customer", "type": "String", "required": "Yes", "desc": "Customer ID"},
        {"name": "paid_amount", "type": "Float", "required": "Yes", "desc": "Amount being paid"},
        {"name": "mode_of_payment", "type": "String", "required": "Yes", "desc": "Mode of payment (Cash, Bank Transfer, etc.)"},
        {"name": "references", "type": "Array", "required": "No", "desc": "Invoice references for allocation"},
        {"name": "references[].reference_doctype", "type": "String", "required": "Yes", "desc": "Always 'Sales Invoice'"},
        {"name": "references[].reference_name", "type": "String", "required": "Yes", "desc": "Invoice name"},
        {"name": "references[].allocated_amount", "type": "Float", "required": "Yes", "desc": "Amount allocated to this invoice"},
    ],
    request_body={
        "customer": "CUST-0001",
        "paid_amount": 110.00,
        "mode_of_payment": "Cash",
        "references": [
            {
                "reference_doctype": "Sales Invoice",
                "reference_name": "SINV-2026-00032",
                "allocated_amount": 110.00
            }
        ]
    },
    response_body={
        "message": {
            "name": "PE-2026-00015",
            "status": "Draft",
            "paid_amount": 110.00,
            "mode_of_payment": "Cash"
        }
    },
)

doc.add_heading("9.2 Cancel Payment Entry", level=2)
add_endpoint_detail(
    API_PREFIX + "cancel_payment_entry", "POST",
    "Cancel a submitted payment entry.",
    params=[
        {"name": "name", "type": "String", "required": "Yes", "desc": "Payment entry name"},
    ],
    request_body={"name": "PE-2026-00015"},
    response_body={"message": {"name": "PE-2026-00015", "status": "Cancelled"}},
)

doc.add_heading("9.3 Get Receipt List", level=2)
add_endpoint_detail(
    API_PREFIX + "get_receipt_list", "GET",
    "Retrieve a paginated list of payment receipts.",
    params=[
        {"name": "customer", "type": "String", "required": "No", "desc": "Filter by customer"},
        {"name": "page_number", "type": "Integer", "required": "No", "desc": "Page number"},
        {"name": "page_size", "type": "Integer", "required": "No", "desc": "Records per page"},
    ],
    response_body={
        "message": {
            "receipts": [
                {
                    "name": "PE-2026-00015",
                    "customer": "CUST-0001",
                    "customer_name": "ABC Trading",
                    "paid_amount": 110.00,
                    "mode_of_payment": "Cash",
                    "posting_date": "2026-03-16",
                    "status": "Submitted"
                }
            ],
            "total_count": 15
        }
    },
)

doc.add_heading("9.4 Get Available Modes of Payment", level=2)
add_endpoint_detail(
    API_PREFIX + "get_available_modes_of_payment", "GET",
    "Retrieve all available modes of payment configured in the system.",
    response_body={
        "message": [
            {"name": "Cash", "type": "Cash"},
            {"name": "Bank Transfer", "type": "Bank"},
            {"name": "Credit Card", "type": "Bank"},
            {"name": "Cheque", "type": "Bank"}
        ]
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 10: CONFIG & SETTINGS
# ══════════════════════════════════════════════════════════════════════════════

section_heading(10, "Config & Settings")

cfg_endpoints = [
    {"path": API_PREFIX + "get_salesperson_config", "method": "GET", "desc": "Get salesperson configuration (warehouse, price list, etc.)"},
    {"path": API_PREFIX + "get_payment_type", "method": "GET", "desc": "Get available payment types"},
    {"path": API_PREFIX + "get_warehouse_list", "method": "GET", "desc": "List all warehouses"},
    {"path": API_PREFIX + "get_price_list", "method": "GET", "desc": "List all price lists"},
    {"path": API_PREFIX + "get_customer_group_list", "method": "GET", "desc": "List all customer groups"},
    {"path": API_PREFIX + "get_available_batches", "method": "GET", "desc": "Get available batches for an item"},
    {"path": API_PREFIX + "sync_items_for_pos", "method": "GET", "desc": "Sync items for POS/offline use"},
    {"path": API_PREFIX + "get_customer_price_list", "method": "GET", "desc": "Get customer-specific price list"},
]
add_endpoint_summary_table(cfg_endpoints)

doc.add_heading("10.1 Get Salesperson Config", level=2)
add_endpoint_detail(
    API_PREFIX + "get_salesperson_config", "GET",
    "Retrieve the salesperson's configuration including default warehouse, price list, and permissions.",
    response_body={
        "message": {
            "sales_person": "SP-001",
            "sales_person_name": "John Doe",
            "default_warehouse": "Van 1 - MT",
            "default_price_list": "Standard Selling",
            "company": "My Trading LLC",
            "territory": "Dubai",
            "allowed_warehouses": ["Van 1 - MT", "Main Store - MT"]
        }
    },
)

doc.add_heading("10.2 Get Available Batches", level=2)
add_endpoint_detail(
    API_PREFIX + "get_available_batches", "GET",
    "Get available batches for a batch-tracked item.",
    params=[
        {"name": "item_code", "type": "String", "required": "Yes", "desc": "Item code"},
        {"name": "warehouse", "type": "String", "required": "No", "desc": "Warehouse filter"},
    ],
    response_body={
        "message": [
            {"batch_no": "BATCH-001", "expiry_date": "2027-01-15", "qty": 500},
            {"batch_no": "BATCH-002", "expiry_date": "2027-06-30", "qty": 300}
        ]
    },
)

doc.add_heading("10.3 Other Config Endpoints", level=2)
doc.add_paragraph("The following endpoints return simple lists and require no parameters (unless noted).")
make_table(
    ["Endpoint", "Method", "Parameters", "Returns"],
    [
        [API_PREFIX + "get_payment_type", "GET", "None", "List of payment types (Cash, Credit, etc.)"],
        [API_PREFIX + "get_warehouse_list", "GET", "None", "List of warehouses with names"],
        [API_PREFIX + "get_price_list", "GET", "None", "List of price lists"],
        [API_PREFIX + "get_customer_group_list", "GET", "None", "List of customer groups"],
        [API_PREFIX + "sync_items_for_pos", "GET", "None", "Full item catalog for offline sync"],
        [API_PREFIX + "get_customer_price_list", "GET", "customer (String)", "Customer-specific price list"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 11: PRINT
# ══════════════════════════════════════════════════════════════════════════════

section_heading(11, "Print")

print_endpoints = [
    {"path": API_PREFIX + "print_document", "method": "GET", "desc": "Get printable HTML for a document"},
    {"path": API_PREFIX + "get_available_print_formats", "method": "GET", "desc": "List available print formats for a doctype"},
    {"path": API_PREFIX + "download_invoice_pdf", "method": "GET", "desc": "Download PDF of an invoice"},
]
add_endpoint_summary_table(print_endpoints)

doc.add_heading("11.1 Print Document", level=2)
add_endpoint_detail(
    API_PREFIX + "print_document", "GET",
    "Generate printable HTML output for any document.",
    params=[
        {"name": "doctype", "type": "String", "required": "Yes", "desc": "Document type (Sales Invoice, Delivery Note, etc.)"},
        {"name": "name", "type": "String", "required": "Yes", "desc": "Document name"},
        {"name": "print_format", "type": "String", "required": "No", "desc": "Print format name"},
    ],
    response_body={"message": "<html>...printable content...</html>"},
)

doc.add_heading("11.2 Get Available Print Formats", level=2)
add_endpoint_detail(
    API_PREFIX + "get_available_print_formats", "GET",
    "List available print formats for a document type.",
    params=[
        {"name": "doctype", "type": "String", "required": "Yes", "desc": "Document type"},
    ],
    response_body={
        "message": [
            {"name": "Standard", "default": True},
            {"name": "POS Invoice", "default": False},
            {"name": "Thermal Print", "default": False}
        ]
    },
)

doc.add_heading("11.3 Download Invoice PDF", level=2)
add_endpoint_detail(
    API_PREFIX + "download_invoice_pdf", "GET",
    "Download a PDF version of an invoice.",
    params=[
        {"name": "name", "type": "String", "required": "Yes", "desc": "Invoice name"},
        {"name": "print_format", "type": "String", "required": "No", "desc": "Print format name"},
    ],
    response_body={"message": "Returns PDF file as binary download"},
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 12: SHIFT MANAGEMENT
# ══════════════════════════════════════════════════════════════════════════════

section_heading(12, "Shift Management")

shift_endpoints = [
    {"path": API_PREFIX + "get_open_shift", "method": "GET", "desc": "Check if there is an open shift for the user"},
    {"path": API_PREFIX + "create_shift_opening_entry", "method": "POST", "desc": "Open a new shift with opening balance"},
    {"path": API_PREFIX + "create_shift_closing_entry", "method": "POST", "desc": "Close the current shift with closing details"},
]
add_endpoint_summary_table(shift_endpoints)

doc.add_heading("12.1 Get Open Shift", level=2)
add_endpoint_detail(
    API_PREFIX + "get_open_shift", "GET",
    "Check whether the current user has an open (active) shift.",
    response_body={
        "message": {
            "has_open_shift": True,
            "shift_name": "SHIFT-2026-00010",
            "opening_date": "2026-03-16",
            "opening_time": "08:00:00",
            "opening_balance": 500.00
        }
    },
)

doc.add_heading("12.2 Create Shift Opening Entry", level=2)
add_endpoint_detail(
    API_PREFIX + "create_shift_opening_entry", "POST",
    "Open a new shift for the salesperson with the cash opening balance.",
    params=[
        {"name": "opening_balance", "type": "Float", "required": "Yes", "desc": "Cash opening balance"},
        {"name": "posting_date", "type": "Date", "required": "No", "desc": "Shift date (defaults to today)"},
    ],
    request_body={
        "opening_balance": 500.00,
        "posting_date": "2026-03-16"
    },
    response_body={
        "message": {
            "name": "SHIFT-2026-00010",
            "status": "Open",
            "opening_balance": 500.00,
            "opening_date": "2026-03-16"
        }
    },
)

doc.add_heading("12.3 Create Shift Closing Entry", level=2)
add_endpoint_detail(
    API_PREFIX + "create_shift_closing_entry", "POST",
    "Close the current open shift with closing balance and payment mode breakdown.",
    params=[
        {"name": "shift_name", "type": "String", "required": "Yes", "desc": "Open shift name"},
        {"name": "closing_balance", "type": "Float", "required": "Yes", "desc": "Cash closing balance"},
    ],
    request_body={
        "shift_name": "SHIFT-2026-00010",
        "closing_balance": 1200.00
    },
    response_body={
        "message": {
            "name": "SHIFT-2026-00010",
            "status": "Closed",
            "opening_balance": 500.00,
            "closing_balance": 1200.00,
            "total_sales": 3500.00,
            "total_collections": 3200.00,
            "closing_date": "2026-03-16",
            "closing_time": "18:00:00"
        }
    },
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX A: ERROR RESPONSE FORMAT
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix A: Standard Error Response Format", level=1)

doc.add_paragraph(
    "All API endpoints return errors in a consistent format. "
    "The HTTP status code indicates the general category of the error, "
    "and the response body provides details."
)

doc.add_heading("Error Response Structure", level=2)
add_json_block("Error Response:", {
    "exc_type": "ValidationError",
    "exception": "frappe.exceptions.ValidationError: Customer is mandatory",
    "_server_messages": "[\"Customer is mandatory\"]"
})

doc.add_heading("Authentication Error", level=2)
add_json_block("403 Forbidden (Missing CSRF Token):", {
    "exc_type": "CSRFTokenError",
    "_server_messages": "[\"CSRF token missing or invalid\"]"
})

doc.add_heading("Session Expired", level=2)
add_json_block("403 Forbidden (Session Expired):", {
    "exc_type": "AuthenticationError",
    "session_expired": 1,
    "_server_messages": "[\"Session Expired. Please login again.\"]"
})

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX B: HTTP STATUS CODES
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix B: HTTP Status Codes", level=1)

make_table(
    ["Status Code", "Meaning", "When It Occurs"],
    [
        ["200", "OK", "Request processed successfully"],
        ["400", "Bad Request", "Invalid parameters or missing required fields"],
        ["401", "Unauthorized", "Invalid credentials during login"],
        ["403", "Forbidden", "Missing/invalid session or CSRF token"],
        ["404", "Not Found", "Endpoint or resource does not exist"],
        ["409", "Conflict", "Document workflow conflict (e.g., submitting already submitted doc)"],
        ["417", "Expectation Failed", "Server-side validation error"],
        ["500", "Internal Server Error", "Unexpected server error"],
        ["502", "Bad Gateway", "Server temporarily unavailable"],
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# APPENDIX C: WORKFLOW DIAGRAMS
# ══════════════════════════════════════════════════════════════════════════════

doc.add_page_break()
doc.add_heading("Appendix C: Workflow Diagrams", level=1)

doc.add_heading("Standard Sales Workflow", level=2)
doc.add_paragraph(
    "The typical van sales workflow follows this sequence:"
)

workflow1 = doc.add_paragraph()
workflow1.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow_text = (
    "Sales Order (SO)\n"
    "       |\n"
    "       v\n"
    "Delivery Note (DN)\n"
    "       |\n"
    "       v\n"
    "Sales Invoice (SI)\n"
    "       |\n"
    "       v\n"
    "Payment Receipt (PE)"
)
run_flow = workflow1.add_run(flow_text)
run_flow.font.name = "Consolas"
run_flow.font.size = Pt(11)
run_flow.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()
doc.add_paragraph("API Call Sequence:")
make_table(
    ["Step", "Action", "API Endpoint", "Method"],
    [
        ["1", "Create Sales Order", API_PREFIX + "create_sales_order", "POST"],
        ["2", "Submit Sales Order", API_PREFIX + "submit_sales_order", "POST"],
        ["3", "Create DN from SO", API_PREFIX + "create_dn_from_so", "POST"],
        ["4", "Submit Delivery Note", API_PREFIX + "submit_delivery_note", "POST"],
        ["5", "Create SI from DN", API_PREFIX + "create_si_from_dn", "POST"],
        ["6", "Submit Sales Invoice", API_PREFIX + "submit_sales_invoice", "POST"],
        ["7", "Create Payment Entry", API_PREFIX + "create_payment_entry", "POST"],
    ],
)

doc.add_paragraph()
doc.add_heading("Direct Sales Workflow (No SO/DN)", level=2)
doc.add_paragraph("For cash sales or direct invoicing:")

workflow2 = doc.add_paragraph()
workflow2.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow_text2 = (
    "Sales Invoice (SI)\n"
    "       |\n"
    "       v\n"
    "Payment Receipt (PE)"
)
run_flow2 = workflow2.add_run(flow_text2)
run_flow2.font.name = "Consolas"
run_flow2.font.size = Pt(11)
run_flow2.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()
doc.add_heading("Returns Workflow", level=2)
doc.add_paragraph("When a customer returns goods:")

workflow3 = doc.add_paragraph()
workflow3.alignment = WD_ALIGN_PARAGRAPH.CENTER
flow_text3 = (
    "Sales Invoice (SI)\n"
    "       |\n"
    "       v\n"
    "Sales Return (Credit Note)\n"
    "       |\n"
    "       v\n"
    "Payment / Adjustment (PE)"
)
run_flow3 = workflow3.add_run(flow_text3)
run_flow3.font.name = "Consolas"
run_flow3.font.size = Pt(11)
run_flow3.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)

doc.add_paragraph()
doc.add_paragraph("API Call Sequence:")
make_table(
    ["Step", "Action", "API Endpoint", "Method"],
    [
        ["1", "Get invoices for return", API_PREFIX + "get_invoices_for_return", "GET"],
        ["2", "Create return (credit note)", API_PREFIX + "create_sales_invoice_return", "POST"],
        ["3", "Submit return invoice", API_PREFIX + "submit_sales_invoice", "POST"],
        ["4", "Create payment/adjustment", API_PREFIX + "create_payment_entry", "POST"],
    ],
)

doc.add_paragraph()
doc.add_heading("SO to SI Conversion Workflow", level=2)
doc.add_paragraph("Convert sales orders directly to invoices (skipping delivery note):")

make_table(
    ["Step", "Action", "API Endpoint", "Method"],
    [
        ["1", "Create Sales Order", API_PREFIX + "create_sales_order", "POST"],
        ["2", "Submit Sales Order", API_PREFIX + "submit_sales_order", "POST"],
        ["3", "Convert SO to SI", API_PREFIX + "convert_so_to_si", "POST"],
        ["4", "Submit Sales Invoice", API_PREFIX + "submit_sales_invoice", "POST"],
    ],
)

# ── SAVE ─────────────────────────────────────────────────────────────────────

output_path = "/Users/sammishthundiyil/frappe-bench/apps/basket4me_pwa/Basket4Me_PWA_API_Reference_V2.docx"
doc.save(output_path)
print(f"Document saved to: {output_path}")
print(f"File size: {os.path.getsize(output_path) / 1024:.1f} KB")
